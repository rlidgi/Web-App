from flask import Flask, request, render_template, redirect, url_for, session, flash, send_file, jsonify, Response, make_response
from io import BytesIO
import PyPDF2
import pdfplumber
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import mammoth
import logging
 # ...existing code...
import pdfplumber
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import mammoth
import logging

# Import analytics module for tracking Facebook ad performance
from analytics import analytics

# Import newsletter system
from newsletter import NewsletterManager, NewsletterConfig

from google_auth_oauthlib.flow import Flow
import google.auth.transport.requests
import google.oauth2.credentials
import google.oauth2.id_token

from flask_dance.contrib.facebook import make_facebook_blueprint, facebook
from flask_login import LoginManager, login_required, login_user, logout_user, UserMixin, current_user
from dotenv import load_dotenv
from datetime import datetime, timezone, timedelta
from typing import Optional
import openai
import os
import json
from docx import Document
import stripe
from azure.data.tables import TableServiceClient, UpdateMode
from azure.identity import DefaultAzureCredential
from urllib.parse import urlparse, urljoin
from urllib.parse import urlencode

app = Flask(__name__)


#####################
# ---- Make `current_user` available in all Jinja templates ----
try:
    from flask_login import LoginManager, current_user as flask_login_current_user, AnonymousUserMixin

    login_manager = LoginManager()
    login_manager.init_app(app)

    @app.context_processor
    def inject_current_user():
        # This exposes the real Flask-Login current_user to Jinja
        return dict(current_user=flask_login_current_user)

except Exception:
    # Flask-Login not installed/configured: provide a safe dummy
    class _Anon:
        is_authenticated = False
    @app.context_processor
    def inject_current_user():
        return dict(current_user=_Anon())



################################





# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()
app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET_KEY') or 'a-very-secret-random-key'


def _is_safe_next_url(target: str) -> bool:
    """Allow only same-host redirects (prefer relative paths) to avoid open redirects."""
    try:
        if not target:
            return False
        # Disallow scheme-relative URLs
        if target.startswith("//"):
            return False
        # Relative path is OK
        if target.startswith("/"):
            return True
        # Otherwise require same host
        ref_url = urlparse(request.host_url)
        test_url = urlparse(urljoin(request.host_url, target))
        return test_url.scheme in ("http", "https") and ref_url.netloc == test_url.netloc
    except Exception:
        return False


def _set_auth_next_from_request() -> None:
    """Capture ?next=... into session for use after login/oauth completes."""
    try:
        nxt = request.args.get("next", "").strip()
        if nxt and _is_safe_next_url(nxt):
            session["auth_next"] = nxt
    except Exception:
        pass


def _pop_auth_next() -> Optional[str]:
    """Pop a safe next URL from session, if any."""
    try:
        nxt = session.pop("auth_next", None)
        if nxt and _is_safe_next_url(nxt):
            return nxt
    except Exception:
        pass
    return None

# Ensure HTTPS URLs in sitemap and external links
app.config['PREFERRED_URL_SCHEME'] = 'https'

# File upload configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'temp_uploads'

# Create upload folder if it doesn't exist
import os
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])



def _append_google_signup_csv(user_id: str, name: str, email: str, created_at_iso: str) -> None:
    """Append a record of a new Google signup to google_signups.csv.

    Columns: timestamp_iso, user_id, name, email
    """
    try:
        import csv
        filename = 'google_signups.csv'
        file_exists = os.path.exists(filename)
        with open(filename, 'a', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if not file_exists:
                writer.writerow(['timestamp_iso', 'user_id', 'name', 'email'])
            writer.writerow([created_at_iso or datetime.now(timezone.utc).isoformat(), user_id or '', name or '', (email or '').strip().lower()])
    except Exception as e:
        logger.error(f"Failed to append google signup CSV: {str(e)}")



@app.before_request
def before_request():
    # Only redirect in production (not in debug mode)
    if not app.debug:
        forwarded_proto = request.headers.get('X-Forwarded-Proto', '').lower()
        app.logger.info(f"[HTTPS REDIRECT] scheme={request.scheme}, X-Forwarded-Proto={forwarded_proto}, port={request.environ.get('SERVER_PORT')}, url={request.url}")
        # Only redirect if X-Forwarded-Proto is present and is 'http'
        if forwarded_proto == 'http':
            app.logger.info("[HTTPS REDIRECT] Redirecting to HTTPS due to X-Forwarded-Proto == 'http'")
            url = request.url.replace('http://', 'https://', 1)
            return redirect(url, code=301)
        # If header is missing, only redirect if scheme is http and port is 80 (default HTTP)
        elif not forwarded_proto and request.scheme == 'http' and request.environ.get('SERVER_PORT') == '80':
            app.logger.info("[HTTPS REDIRECT] Redirecting to HTTPS due to scheme == 'http' and port == 80")
            url = request.url.replace('http://', 'https://', 1)
            return redirect(url, code=301)
        # Otherwise, do not redirect (prevents loop)

# Add security headers for HTTPS enforcement
@app.before_request
def before_request():
    if not app.debug:
        forwarded_proto = request.headers.get('X-Forwarded-Proto', '').lower()
        app.logger.info(f"[HTTPS REDIRECT] scheme={request.scheme}, X-Forwarded-Proto={forwarded_proto}, port={request.environ.get('SERVER_PORT')}, url={request.url}")
        # Only redirect if X-Forwarded-Proto is present and is 'http'
        if forwarded_proto == 'http':
            app.logger.info("[HTTPS REDIRECT] Redirecting to HTTPS due to X-Forwarded-Proto == 'http'")
            url = request.url.replace('http://', 'https://', 1)
            return redirect(url, code=301)
        # If header is missing, only redirect if scheme is http and port is 80 (default HTTP)
        elif not forwarded_proto and request.scheme == 'http' and request.environ.get('SERVER_PORT') == '80':
            app.logger.info("[HTTPS REDIRECT] Redirecting to HTTPS due to scheme == 'http' and port == 80")
            url = request.url.replace('http://', 'https://', 1)
            return redirect(url, code=301)
        # Otherwise, do not redirect (prevents loop)



# Google OAuth Configuration
GOOGLE_CLIENT_SECRET_FILE = os.path.join(os.path.dirname(__file__), "client_secret.json")

# Check if we have environment variables for Google OAuth (production)
if os.getenv('GOOGLE_CLIENT_ID') and os.getenv('GOOGLE_CLIENT_SECRET'):
    # Use environment variables (more secure for production)
    GOOGLE_CLIENT_ID = os.getenv('GOOGLE_CLIENT_ID')
    GOOGLE_CLIENT_SECRET = os.getenv('GOOGLE_CLIENT_SECRET')
    USE_ENV_CREDENTIALS = True
else:
    # Use file-based credentials (for development)
    USE_ENV_CREDENTIALS = False

GOOGLE_SCOPES = [
    "https://www.googleapis.com/auth/userinfo.email",
    "https://www.googleapis.com/auth/userinfo.profile", 
    "openid"
]


facebook_bp = make_facebook_blueprint(
    client_id=os.getenv("FACEBOOK_APP_ID"),
    client_secret=os.getenv("FACEBOOK_APP_SECRET"),
    scope="email",
    redirect_to="facebook_callback",  # must be HTTPS
)


app.register_blueprint(facebook_bp, url_prefix="/login")  # MUST be before the route below


# Flask-Login
login_manager = LoginManager(app)
login_manager.login_view = "login"

# Define a list of admin email addresses
ADMIN_EMAILS = ["yaronyaronlid@gmail.com"]

# User storage file
USERS_FILE = "users_data.json"

# Password reset tokens storage
RESET_TOKENS_FILE = "reset_tokens.json"
RESET_TOKEN_EXPIRY_HOURS = 24  # Tokens expire after 24 hours



class User(UserMixin):
    def __init__(self, id, name, email, password_hash=None, is_new=False, created_at=None):
        self.id = id
        self.name = name
        self.email = email
        self.password_hash = password_hash
        self.is_new = is_new
        self.created_at = created_at or datetime.now(timezone.utc).isoformat()
        # Determine if the user is an admin based on their email
        self.is_admin = email in ADMIN_EMAILS
    
    def set_password(self, password):
        """Hash and set password"""
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        """Check if provided password matches hash"""
        if not self.password_hash:
            return False
        return check_password_hash(self.password_hash, password)
    
    def to_dict(self):
        """Convert user to dictionary for JSON storage"""
        return {
            'id': self.id,
            'name': self.name,
            'email': self.email,
            'password_hash': self.password_hash,
            'created_at': self.created_at,
            'is_admin': self.is_admin
        }
    
    @classmethod
    def from_dict(cls, data):
        """Create user from dictionary"""
        user = cls(
            id=data['id'],
            name=data['name'],
            email=data['email'],
            password_hash=data.get('password_hash'),
            created_at=data.get('created_at')
        )
        return user

def load_users():
    """Load users from JSON file"""
    try:
        if os.path.exists(USERS_FILE):
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                users_data = json.load(f)
                return {user_id: User.from_dict(user_data) for user_id, user_data in users_data.items()}
        return {}
    except Exception as e:
        print(f"Error loading users: {e}")
        return {}

def save_users():
    """Save users to JSON file"""
    try:
        users_data = {user_id: user.to_dict() for user_id, user in users.items()}
        with open(USERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(users_data, f, indent=2, ensure_ascii=False)
        print(f"✅ Saved {len(users)} users to persistent storage")
    except Exception as e:
        print(f"Error saving users: {e}")



def add_user(user):
    """Add a user and save to persistent storage"""
    users[user.id] = user
    save_users()
    print(f"✅ Added user: {user.name} ({user.email})")

# Load existing users on startup
users = load_users()
print(f"📊 Loaded {len(users)} users from persistent storage")



@login_manager.user_loader
def load_user(user_id):
    return users.get(user_id)

@app.route("/login", methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        # If already logged in and a next is provided, honor it.
        _set_auth_next_from_request()
        nxt = _pop_auth_next()
        return redirect(nxt or url_for('index'))

    # Clear any lingering flash messages
    session.pop('_flashes', None)

    # Capture intended return URL for post-auth redirect
    _set_auth_next_from_request()

    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'login':
            # Handle email/password login
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')
            
            if not email or not password:
                flash('Please enter both email and password.', 'danger')
                return render_template("login.html")
            
            # Find user by email
            user = None
            for user_id, u in users.items():
                if u.email.lower() == email:
                    user = u
                    break
            
            if user and user.password_hash and user.check_password(password):
                login_user(user)
                # Save pending revision if it exists
                pending = session.pop('pending_revision', None)
                if pending:
                    import uuid
                    try:
                        save_resume_revision(
                            user.id,
                            str(uuid.uuid4()),
                            pending['revised_resume'],
                            feedback=pending.get('feedback'),
                            original_resume=pending.get('original_resume'),
                            job_description=pending.get('job_description')
                        )
                    except FreeTierLimitReached:
                        flash("You've reached the free tier limit (2 revisions). Upgrade to save unlimited revisions.", "danger")
                    except Exception:
                        pass
                flash('You have been successfully logged in!', 'success')
                nxt = _pop_auth_next()
                return redirect(nxt or url_for('my_revisions'))
            else:
                flash('Invalid email or password.', 'danger')
                return render_template("login.html")
        
        elif action == 'register':
            # Handle email/password registration
            name = request.form.get('name', '').strip()
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')
            confirm_password = request.form.get('confirm_password', '')
            
            # Validation
            if not name or not email or not password:
                flash('Please fill in all fields.', 'danger')
                return render_template("login.html")
            
            if len(password) < 8:
                flash('Password must be at least 8 characters long.', 'danger')
                return render_template("login.html")
            
            if password != confirm_password:
                flash('Passwords do not match.', 'danger')
                return render_template("login.html")
            
            # Check if email already exists
            for user_id, u in users.items():
                if u.email.lower() == email:
                    flash('An account with this email already exists. Please login instead.', 'danger')
                    return render_template("login.html")
            
            # Create new user
            import uuid
            user_id = f"email_{uuid.uuid4().hex[:16]}"
            user = User(user_id, name, email, is_new=True)
            user.set_password(password)
            add_user(user)
            
            # Persist profile to Azure Users table
            try:
                upsert_user_profile_azure(user)
            except Exception:
                pass
            
            login_user(user)
            flash('Account created successfully! Welcome to ResumaticAI!', 'success')
            
            # Save pending revision if it exists
            pending = session.pop('pending_revision', None)
            if pending:
                try:
                    save_resume_revision(
                        user.id,
                        str(uuid.uuid4()),
                        pending['revised_resume'],
                        feedback=pending.get('feedback'),
                        original_resume=pending.get('original_resume'),
                        job_description=pending.get('job_description')
                    )
                except FreeTierLimitReached:
                    flash("You've reached the free tier limit (2 revisions). Upgrade to save unlimited revisions.", "danger")
                except Exception:
                    pass
            
            nxt = _pop_auth_next()
            return redirect(nxt or url_for('my_revisions'))

    return render_template("login.html")

# Password Reset Functions
def load_reset_tokens():
    """Load reset tokens from JSON file"""
    try:
        if os.path.exists(RESET_TOKENS_FILE):
            with open(RESET_TOKENS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    except Exception as e:
        logger.error(f"Error loading reset tokens: {e}")
        return {}

def save_reset_tokens(tokens):
    """Save reset tokens to JSON file"""
    try:
        with open(RESET_TOKENS_FILE, 'w', encoding='utf-8') as f:
            json.dump(tokens, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logger.error(f"Error saving reset tokens: {e}")

def generate_reset_token():
    """Generate a secure random token for password reset"""
    import secrets
    return secrets.token_urlsafe(32)

def send_password_reset_email(email, token, user_name):
    """Send password reset email to user"""
    try:
        _load_email_config_if_missing()
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        
        smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
        smtp_port = int(os.getenv('SMTP_PORT', '587'))
        auth_email = os.getenv('NEWSLETTER_EMAIL', '').strip()
        auth_password = os.getenv('NEWSLETTER_PASSWORD', '').strip().replace(' ', '')
        
        if not auth_email or not auth_password:
            raise ValueError('Email credentials not configured. Set NEWSLETTER_EMAIL and NEWSLETTER_PASSWORD.')
        
        # Generate reset URL
        reset_url = url_for('reset_password', token=token, _external=True)
        
        # Create email
        subject = 'Reset Your Password - ResumaticAI'
        html_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2 style="color: #2563eb;">Reset Your Password</h2>
                <p>Hello {user_name},</p>
                <p>We received a request to reset your password for your ResumaticAI account.</p>
                <p>Click the button below to reset your password:</p>
                <div style="margin: 30px 0;">
                    <a href="{reset_url}" style="background-color: #2563eb; color: white; padding: 12px 24px; text-decoration: none; border-radius: 5px; display: inline-block;">Reset Password</a>
                </div>
                <p>Or copy and paste this link into your browser:</p>
                <p style="word-break: break-all; color: #666;">{reset_url}</p>
                <p>This link will expire in 24 hours.</p>
                <p>If you didn't request a password reset, please ignore this email.</p>
                <hr style="border: none; border-top: 1px solid #eee; margin: 20px 0;">
                <p style="color: #666; font-size: 12px;">ResumaticAI Team</p>
            </div>
        </body>
        </html>
        """
        
        text_body = f"""
Reset Your Password - ResumaticAI

Hello {user_name},

We received a request to reset your password for your ResumaticAI account.

Click the link below to reset your password:
{reset_url}

This link will expire in 24 hours.

If you didn't request a password reset, please ignore this email.

ResumaticAI Team
        """
        
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = f"ResumaticAI <{auth_email}>"
        msg['To'] = email
        
        msg.attach(MIMEText(text_body, 'plain'))
        msg.attach(MIMEText(html_body, 'html'))
        
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(auth_email, auth_password)
        server.send_message(msg)
        server.quit()
        
        logger.info(f"Password reset email sent to {email}")
        return True
    except Exception as e:
        logger.error(f"Error sending password reset email: {str(e)}")
        return False

@app.route("/forgot-password", methods=['GET', 'POST'])
def forgot_password():
    """Handle forgot password requests"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        
        if not email:
            flash('Please enter your email address.', 'danger')
            return render_template("forgot_password.html")
        
        # Find user by email
        user = None
        for user_id, u in users.items():
            if u.email.lower() == email:
                user = u
                break
        
        # Always show success message (security: don't reveal if email exists)
        flash('If an account exists with that email, a password reset link has been sent.', 'success')
        
        if user and user.password_hash:  # Only send if user has password (email signup)
            # Generate reset token
            token = generate_reset_token()
            expiry_time = datetime.now(timezone.utc).timestamp() + (RESET_TOKEN_EXPIRY_HOURS * 3600)
            
            # Save token
            tokens = load_reset_tokens()
            tokens[token] = {
                'user_id': user.id,
                'email': user.email,
                'expires_at': expiry_time
            }
            save_reset_tokens(tokens)
            
            # Send email
            send_password_reset_email(user.email, token, user.name)
        
        return redirect(url_for('login'))
    
    return render_template("forgot_password.html")

@app.route("/reset-password/<token>", methods=['GET', 'POST'])
def reset_password(token):
    """Handle password reset with token"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    # Load tokens
    tokens = load_reset_tokens()
    
    if token not in tokens:
        flash('Invalid or expired reset link. Please request a new password reset.', 'danger')
        return redirect(url_for('forgot_password'))
    
    token_data = tokens[token]
    current_time = datetime.now(timezone.utc).timestamp()
    
    # Check if token expired
    if current_time > token_data['expires_at']:
        # Remove expired token
        del tokens[token]
        save_reset_tokens(tokens)
        flash('This reset link has expired. Please request a new password reset.', 'danger')
        return redirect(url_for('forgot_password'))
    
    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        
        # Validation
        if not password or not confirm_password:
            flash('Please fill in both password fields.', 'danger')
            return render_template("reset_password.html", token=token, valid=True)
        
        if len(password) < 8:
            flash('Password must be at least 8 characters long.', 'danger')
            return render_template("reset_password.html", token=token, valid=True)
        
        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return render_template("reset_password.html", token=token, valid=True)
        
        # Get user and update password
        user_id = token_data['user_id']
        user = users.get(user_id)
        
        if user:
            user.set_password(password)
            add_user(user)  # Save updated password
            
            # Remove used token
            del tokens[token]
            save_reset_tokens(tokens)
            
            flash('Your password has been reset successfully! Please login with your new password.', 'success')
            return redirect(url_for('login'))
        else:
            flash('User not found. Please contact support.', 'danger')
            return redirect(url_for('login'))
    
    return render_template("reset_password.html", token=token, valid=True)

@app.route("/logout")
@login_required
def logout():
    # Clear any OAuth tokens or state
    session.pop('state', None)
    session.pop('google_oauth_token', None)
    session.pop('facebook_oauth_token', None)
    session.pop('pending_revision', None)
    # Clear full session
    session.clear()
    # Log out the user
    logout_user()
    flash('You have been successfully logged out.', 'success')
    return redirect(url_for('index'))


@app.route("/login/google")
def google_login():
    # Clear any lingering flash messages
    session.pop('_flashes', None)

    if current_user.is_authenticated:
        return redirect(url_for('index'))

    # Capture intended return URL for post-auth redirect
    _set_auth_next_from_request()
    
    if USE_ENV_CREDENTIALS:
        # Use environment variables for credentials
        from google_auth_oauthlib.flow import Flow
        client_config = {
            "web": {
                "client_id": GOOGLE_CLIENT_ID,
                "client_secret": GOOGLE_CLIENT_SECRET,
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
                "redirect_uris": [url_for("google_callback", _external=True)]
            }
        }
        flow = Flow.from_client_config(
            client_config,
            scopes=GOOGLE_SCOPES,
            redirect_uri=url_for("google_callback", _external=True)
        )
    else:
        # Use file-based credentials
        flow = Flow.from_client_secrets_file(
            GOOGLE_CLIENT_SECRET_FILE,
            scopes=GOOGLE_SCOPES,
            redirect_uri=url_for("google_callback", _external=True)
        )
    
    auth_url, state = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true",
        prompt="select_account"  # 👈 This line is important
    )
    session["state"] = state
    return redirect(auth_url)





# Google OAuth Callback
@app.route("/login/google/authorized")
def google_callback():
    if current_user.is_authenticated:
        nxt = _pop_auth_next()
        return redirect(nxt or url_for('my_revisions'))
    
    if USE_ENV_CREDENTIALS:
        # Use environment variables for credentials
        client_config = {
            "web": {
                "client_id": GOOGLE_CLIENT_ID,
                "client_secret": GOOGLE_CLIENT_SECRET,
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
                "redirect_uris": [url_for("google_callback", _external=True)]
            }
        }
        flow = Flow.from_client_config(
            client_config,
            scopes=GOOGLE_SCOPES,
            state=session["state"],
            redirect_uri=url_for("google_callback", _external=True)
        )
    else:
        # Use file-based credentials
        flow = Flow.from_client_secrets_file(
            GOOGLE_CLIENT_SECRET_FILE,
            scopes=GOOGLE_SCOPES,
            state=session["state"],
            redirect_uri=url_for("google_callback", _external=True)
        )

    flow.fetch_token(authorization_response=request.url)
    credentials = flow.credentials
    
    # Add clock skew tolerance for token verification with retry logic
    import time
    max_retries = 3
    for attempt in range(max_retries):
        try:
            user_info = google.oauth2.id_token.verify_oauth2_token(
                credentials.id_token, 
                google.auth.transport.requests.Request(),
                clock_skew_in_seconds=30  # Allow 30 seconds clock skew
            )
            break  # Success, exit retry loop
        except google.auth.exceptions.InvalidValue as e:
            if "Token used too early" in str(e) and attempt < max_retries - 1:
                # Wait 2 seconds and retry
                time.sleep(2)
                continue
            else:
                raise  # Re-raise if not a timing issue or max retries reached
    user_id = user_info["sub"]
    is_new = user_id not in users
    user = User(user_id, user_info["name"], user_info.get("email", ""), is_new=is_new)
    add_user(user)  # Use add_user to save persistently
    # Record only first-time Google signups
    try:
        if is_new and str(user_id).isdigit():
            _append_google_signup_csv(user.id, user.name, user.email, user.created_at)
    except Exception:
        pass
    # Persist profile to Azure Users table
    try:
        upsert_user_profile_azure(user)
    except Exception:
        pass
    login_user(user)
    # Save pending revision if it exists
    pending = session.pop('pending_revision', None)
    if pending:
        import uuid
        try:
            save_resume_revision(
                user.id,
                str(uuid.uuid4()),
                pending['revised_resume'],
                feedback=pending.get('feedback'),
                original_resume=pending.get('original_resume'),
                job_description=pending.get('job_description')
            )
        except FreeTierLimitReached:
            flash("You've reached the free tier limit (2 revisions). Upgrade to save unlimited revisions.", "danger")
        except Exception:
            pass
    nxt = _pop_auth_next()
    return redirect(nxt or url_for('my_revisions'))
    

@app.route("/login/facebook/authorized")
def facebook_callback():
    if not facebook.authorized:
        flash("Facebook login failed.", "danger")
        return redirect(url_for("login"))
    resp = facebook.get("/me?fields=id,name,email")
    if not resp.ok:
        flash("Failed to fetch Facebook user info.", "danger")
        return redirect(url_for("login"))

    fb_info = resp.json()
    user_id = f"facebook_{fb_info['id']}"
    is_new = user_id not in users
    user = User(user_id, fb_info["name"], fb_info.get("email", ""), is_new=is_new)
    add_user(user)  # Use add_user to save persistently
    # Persist profile to Azure Users table
    try:
        upsert_user_profile_azure(user)
    except Exception:
        pass
    login_user(user)
    nxt = _pop_auth_next()
    return redirect(nxt or url_for("index"))





from openai import OpenAI

def revise_resume(resume_text, job_description=None):
    try:
        # Initialize the client inside the function to avoid blocking startup
        # Explicitly pass API key to handle Azure environment
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set")
        
        # Configure timeout and other parameters for Azure reliability
        client = OpenAI(
            api_key=api_key,
            timeout=60.0,  # 60 second timeout
            max_retries=3  # Retry up to 3 times on transient errors
        )

        # Build the prompt with optional job description
        job_desc_section = ""
        if job_description:
            job_desc_section = f"""
Job Description to tailor the resume towards:
\"\"\"{job_description}\"\"\"
"""

        prompt = f"""You are a resume optimization expert. Analyze the following resume and provide a comprehensive revision.

Resume to analyze:
\"\"\"{resume_text}\"\"\"{job_desc_section}

Provide your response in the following JSON format ONLY (no markdown, no additional text):
{{
  "revised_resume": "The complete revised version of the resume with improved clarity, structure, and professional tone",
  "feedback": {{
    "overall_score": <number between 0-100>,
    "subscores": {{
      "Content": <number between 0-100>,
      "Format": <number between 0-100>,
      "Optimization": <number between 0-100>,
      "Best Practices": <number between 0-100>,
      "Application Readiness": <number between 0-100>
    }},
    "improvement_items": [
      {{
        "category": "<category name>",
        "message": "<improvement suggestion>",
        "severity": "<error|warning|suggestion>",
        "example_lines": "<relevant example from resume>"
      }}
    ]
  }}
}}

Focus on:
1. Improving clarity, structure, and professional tone
2. Enhancing impact of achievements and experiences
3. Optimizing for ATS systems
4. Maintaining consistency in formatting
5. Adding missing elements that would strengthen the resume

Respond with ONLY the JSON object, no markdown formatting or additional text."""

        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}]
            )
        except Exception as e:
            error_msg = str(e)
            print(f"OpenAI API Error: {error_msg}")
            logger.error(f"OpenAI API Error details: {type(e).__name__}: {error_msg}")
            # Re-raise with more details for debugging
            raise ValueError(f"Failed to connect to OpenAI API: {error_msg}")

        raw_content = response.choices[0].message.content
        print("=== GPT RESPONSE ===")
        print(raw_content)

        # Remove markdown backticks if present
        if raw_content.startswith("```"):
            import re
            raw_content = re.sub(r"^```(?:json)?\n", "", raw_content)
            raw_content = re.sub(r"\n```$", "", raw_content)

        try:
            data = json.loads(raw_content)
            if "revised_resume" not in data or "feedback" not in data:
                raise ValueError("Missing required keys in JSON response")
            return data["revised_resume"], data["feedback"]
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {str(e)}")
            print(f"Raw content: {raw_content}")
            raise ValueError("Invalid response format from OpenAI. Please try again.")
        except KeyError as e:
            print(f"Missing key in response: {str(e)}")
            raise ValueError("Incomplete response from OpenAI. Please try again.")

    except Exception as e:
        print("=== ERROR OCCURRED ===")
        import traceback
        traceback.print_exc()
        raise


def revised_resume_formatted(revised_resume):
    
    prompt2 = f"Please format the following resume for better readability:\n\n{revised_resume}"
    response2 = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt2}]
            )
    return response2.choices[0].message.content


def revised_resume_template(revised_resume):
    resume_text = request.form.get("resume", "").strip()
    
    # Validate that we have resume content
    if not resume_text:
        flash("Please upload a resume file or paste your resume content", 'danger')
        return redirect(url_for('index', scroll_to_form='true'))
    
    # Get job description (optional)
    job_description = request.form.get("jobDescription", "").strip()
    
    # Process the resume
    print(f"Resume text preview (first 200 chars): {resume_text[:200]}...")
    print(f"Job description preview: {job_description[:100] if job_description else 'None'}...")
    revised_resume, feedback = revise_resume(resume_text, job_description)

@app.route("/")
def index():
    # Track the visit with source attribution (only if not already tracked)
    if 'visit_tracked' not in session:
        source_info = analytics.track_visit(request)
        session['visit_tracked'] = True
        session['traffic_source'] = source_info
        app.logger.info(f"Homepage visit tracked from server-side: {source_info.get('type', 'unknown')}")
    else:
        # Use existing traffic source from session
        source_info = session.get('traffic_source', {'type': 'organic'})
        app.logger.info("Homepage visit already tracked in session, skipping duplicate")
    
    current_year = datetime.now().year
    scroll_to_form = (request.args.get('scroll_to_form', '').lower() == 'true')
    return render_template("index.html", year=current_year, user=current_user, scroll_to_form=scroll_to_form)


@app.route("/start")
def start():
    current_year = datetime.now().year
    return render_template("start.html", year=current_year, user=current_user)

@app.route("/plans")
def plans():
    current_year = datetime.now().year
    return render_template("plans.html", year=current_year, user=current_user)


def _get_plan_config(plan_id: str) -> Optional[dict]:
    pid = (plan_id or '').strip()
    if not pid:
        return None
    # Plan IDs must match templates/plans.html
    if pid == 'trial_14d':
        return {
            'id': pid,
            'label': '2-Week Trial',
            'price': '$1.85',
            'plan_status': 'trial',
            'duration_days': 14,
        }
    if pid == 'monthly_10_95':
        return {
            'id': pid,
            'label': 'Monthly',
            'price': '$10.95 / month',
            'plan_status': 'monthly',
            'duration_days': 31,
        }
    if pid == 'annual_6_95':
        return {
            'id': pid,
            'label': 'Annual',
            'price': '$6.95 / month (billed annually)',
            'plan_status': 'annual',
            'duration_days': 365,
        }
    return None


def _stripe_enabled() -> bool:
    return bool((os.getenv('STRIPE_SECRET_KEY') or '').strip())


def _get_stripe_customer_id_from_azure(user_id: str) -> str:
    try:
        table_client = get_users_table_client()
        e = table_client.get_entity(partition_key=str(user_id), row_key='profile')
        return str(e.get('stripe_customer_id') or '').strip()
    except Exception:
        return ''


def _get_stripe_subscription_id_from_azure(user_id: str) -> str:
    try:
        table_client = get_users_table_client()
        e = table_client.get_entity(partition_key=str(user_id), row_key='profile')
        return str(e.get('stripe_subscription_id') or '').strip()
    except Exception:
        return ''


def _find_stripe_customer_id_by_email(email: str) -> str:
    """Best-effort lookup for Stripe customer id by email (helps when webhook/profile hasn't saved ids yet)."""
    e = (email or '').strip()
    if not e or not _stripe_enabled():
        return ''
    try:
        stripe.api_key = (os.getenv("STRIPE_SECRET_KEY") or "").strip()
        # Prefer search when available
        try:
            res = stripe.Customer.search(query=f"email:'{e}'", limit=1)
            data = list(getattr(res, 'data', []) or [])
            if data:
                return str(getattr(data[0], 'id', '') or '').strip()
        except Exception:
            pass
        # Fallback: list by email
        res2 = stripe.Customer.list(email=e, limit=1)
        data2 = list(getattr(res2, 'data', []) or [])
        if data2:
            return str(getattr(data2[0], 'id', '') or '').strip()
    except Exception:
        return ''
    return ''


def _find_trialing_subscription_for_customer(customer_id: str) -> Optional[dict]:
    """Return a trialing subscription object (Stripe) for a customer, or None."""
    cid = (customer_id or "").strip()
    if not cid or not _stripe_enabled():
        return None
    try:
        stripe.api_key = (os.getenv("STRIPE_SECRET_KEY") or "").strip()
        res = stripe.Subscription.list(customer=cid, status="trialing", limit=1)
        data = list(getattr(res, "data", []) or [])
        if not data:
            return None
        sub = data[0]
        try:
            sub = stripe.Subscription.retrieve(sub.id, expand=["items.data"])
        except Exception:
            pass
        return sub
    except Exception:
        return None

def _get_paid_until_from_stripe(subscription_id: str) -> str:
    """Return ISO timestamp (UTC) for next renewal/end using Stripe subscription, or '' if unknown."""
    sid = (subscription_id or '').strip()
    if not sid:
        return ''
    if not _stripe_enabled():
        return ''
    try:
        stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
        sub = stripe.Subscription.retrieve(sid)
        # Prefer current_period_end; for trialing subscriptions, trial_end can be useful too.
        trial_end = getattr(sub, 'trial_end', None)
        current_period_end = getattr(sub, 'current_period_end', None)
        ts = current_period_end or trial_end
        if ts:
            return datetime.fromtimestamp(int(ts), tz=timezone.utc).isoformat()
        return ''
    except Exception:
        return ''


def _add_interval_approx(dt: datetime, interval: str, interval_count: int) -> datetime:
    """Approximate interval math without extra deps (good enough for UI labels)."""
    c = int(interval_count or 1)
    if c < 1:
        c = 1
    interval = (interval or '').strip().lower()
    if interval == 'year':
        return dt + timedelta(days=365 * c)
    if interval == 'month':
        return dt + timedelta(days=30 * c)
    if interval == 'week':
        return dt + timedelta(days=7 * c)
    if interval == 'day':
        return dt + timedelta(days=1 * c)
    return dt


def _get_stripe_plan_dates_for_customer(customer_id: str) -> dict:
    """Return best-effort plan date info from Stripe for a customer.

    Output keys:
      - subscription_id
      - status
      - interval_label (e.g. 'Annual'/'Monthly'/'' )
      - next_billing_iso (trial_end if trialing else current_period_end)
      - paid_through_est_iso (if trialing, trial_end + interval; else current_period_end)
    """
    cid = (customer_id or '').strip()
    if not cid or not _stripe_enabled():
        return {}
    try:
        stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
        subs = stripe.Subscription.list(customer=cid, status='all', limit=20)
        data = list(getattr(subs, 'data', []) or [])
        if not data:
            return {}

        def _rank(sub):
            status = str(getattr(sub, 'status', '') or '').lower()
            current_period_end = int(getattr(sub, 'current_period_end', 0) or 0)
            # Prefer active > trialing > others; then later period end.
            status_rank = 0
            if status == 'active':
                status_rank = 3
            elif status == 'trialing':
                status_rank = 2
            elif status in ('past_due', 'unpaid'):
                status_rank = 1
            return (status_rank, current_period_end)

        best = sorted(data, key=_rank, reverse=True)[0]
        # Re-fetch with expanded price/recurring so interval math works reliably (Stripe list results can be "thin")
        try:
            best = stripe.Subscription.retrieve(best.id, expand=["items.data.price"])
        except Exception:
            pass
        status = str(getattr(best, 'status', '') or '')
        trial_end = getattr(best, 'trial_end', None)
        current_period_end = getattr(best, 'current_period_end', None)

        # Interval label (use first item). Stripe objects can sometimes deserialize into plain dicts,
        # so handle both dict and StripeObject attribute access.
        interval = ''
        interval_count = 1
        try:
            items = getattr(best, 'items', None)
            items_data = getattr(items, 'data', []) if items else []
            if items_data:
                price = getattr(items_data[0], 'price', None)
                if isinstance(price, dict):
                    recurring = price.get('recurring')
                else:
                    recurring = getattr(price, 'recurring', None) if price else None
                if isinstance(recurring, dict):
                    interval = str(recurring.get('interval') or '')
                    interval_count = int(recurring.get('interval_count') or 1)
                else:
                    interval = str(getattr(recurring, 'interval', '') or '')
                    interval_count = int(getattr(recurring, 'interval_count', 1) or 1)
        except Exception:
            interval = ''
            interval_count = 1

        interval_label = ''
        if interval == 'year':
            interval_label = 'Annual'
        elif interval == 'month':
            interval_label = 'Monthly'

        next_ts = None
        if str(status).lower() == 'trialing' and trial_end:
            next_ts = int(trial_end)
        elif current_period_end:
            next_ts = int(current_period_end)

        paid_through_est_ts = None
        if str(status).lower() == 'trialing' and trial_end:
            base = datetime.fromtimestamp(int(trial_end), tz=timezone.utc)
            paid_through_est_ts = int(_add_interval_approx(base, interval, interval_count).timestamp())
        elif current_period_end:
            paid_through_est_ts = int(current_period_end)

        def _ts_to_iso(ts: Optional[int]) -> str:
            try:
                if not ts:
                    return ''
                return datetime.fromtimestamp(int(ts), tz=timezone.utc).isoformat()
            except Exception:
                return ''

        return {
            'subscription_id': str(getattr(best, 'id', '') or ''),
            'status': status,
            'interval_label': interval_label,
            'next_billing_iso': _ts_to_iso(next_ts),
            'paid_through_est_iso': _ts_to_iso(paid_through_est_ts),
        }
    except Exception:
        return {}


def _get_stripe_plan_dates_for_subscription(subscription_id: str) -> dict:
    """Fallback when we only have a subscription id (e.g., profile missing stripe_customer_id)."""
    sid = (subscription_id or '').strip()
    if not sid or not _stripe_enabled():
        return {}
    try:
        stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
        sub = stripe.Subscription.retrieve(sid, expand=["items.data.price"])
        status = str(getattr(sub, 'status', '') or '')
        trial_end = getattr(sub, 'trial_end', None)
        current_period_end = getattr(sub, 'current_period_end', None)

        interval = ''
        interval_count = 1
        try:
            items = getattr(sub, 'items', None)
            items_data = getattr(items, 'data', []) if items else []
            if items_data:
                price = getattr(items_data[0], 'price', None)
                if isinstance(price, dict):
                    recurring = price.get('recurring')
                else:
                    recurring = getattr(price, 'recurring', None) if price else None
                if isinstance(recurring, dict):
                    interval = str(recurring.get('interval') or '')
                    interval_count = int(recurring.get('interval_count') or 1)
                else:
                    interval = str(getattr(recurring, 'interval', '') or '')
                    interval_count = int(getattr(recurring, 'interval_count', 1) or 1)
        except Exception:
            interval = ''
            interval_count = 1

        interval_label = ''
        if interval == 'year':
            interval_label = 'Annual'
        elif interval == 'month':
            interval_label = 'Monthly'

        next_ts = None
        if str(status).lower() == 'trialing' and trial_end:
            next_ts = int(trial_end)
        elif current_period_end:
            next_ts = int(current_period_end)

        paid_through_est_ts = None
        if str(status).lower() == 'trialing' and trial_end:
            base = datetime.fromtimestamp(int(trial_end), tz=timezone.utc)
            paid_through_est_ts = int(_add_interval_approx(base, interval, interval_count).timestamp())
        elif current_period_end:
            paid_through_est_ts = int(current_period_end)

        def _ts_to_iso(ts: Optional[int]) -> str:
            try:
                if not ts:
                    return ''
                return datetime.fromtimestamp(int(ts), tz=timezone.utc).isoformat()
            except Exception:
                return ''

        return {
            'subscription_id': str(getattr(sub, 'id', '') or ''),
            'status': status,
            'interval_label': interval_label,
            'next_billing_iso': _ts_to_iso(next_ts),
            'paid_through_est_iso': _ts_to_iso(paid_through_est_ts),
        }
    except Exception:
        return {}

def _get_subscription_price_id_and_recurring(sub_obj) -> tuple[str, str, int]:
    """Return (price_id, interval, interval_count) from a subscription object; robust to dict/StripeObject."""
    try:
        items = getattr(sub_obj, 'items', None)
        items_data = getattr(items, 'data', []) if items else []
        first_item = items_data[0] if items_data else None
        # items_data entries can be StripeObjects or dicts
        if isinstance(first_item, dict):
            price = first_item.get('price')
        else:
            price = getattr(first_item, 'price', None) if first_item else None
        # Sometimes Stripe returns a bare price id string
        if isinstance(price, str):
            return (price.strip(), '', 1)
        if isinstance(price, dict):
            price_id = str(price.get('id') or '').strip()
            recurring = price.get('recurring')
        else:
            price_id = str(getattr(price, 'id', '') or '').strip() if price else ''
            recurring = getattr(price, 'recurring', None) if price else None

        interval = ''
        interval_count = 1
        if isinstance(recurring, dict):
            interval = str(recurring.get('interval') or '').strip()
            interval_count = int(recurring.get('interval_count') or 1)
        else:
            interval = str(getattr(recurring, 'interval', '') or '').strip()
            interval_count = int(getattr(recurring, 'interval_count', 1) or 1)
        return (price_id, interval, interval_count)
    except Exception:
        return ('', '', 1)

def _stripe_obj_get(obj, key: str, default=None):
    """Safely read key from StripeObject or dict (some stripe versions return dict-like objects)."""
    try:
        if obj is None:
            return default
        if isinstance(obj, dict):
            return obj.get(key, default)
        # StripeObject supports .get in many versions
        if hasattr(obj, "get"):
            return obj.get(key, default)
        return getattr(obj, key, default)
    except Exception:
        return default

def _stripe_upcoming_invoice(customer_id: str, subscription_id: str):
    """Get upcoming invoice using a method compatible with older stripe python versions."""
    cid = (customer_id or "").strip()
    sid = (subscription_id or "").strip()
    if not cid or not sid:
        return None
    try:
        # Newer stripe versions have stripe.Invoice.upcoming(...)
        if hasattr(stripe, "Invoice") and hasattr(stripe.Invoice, "upcoming"):
            return stripe.Invoice.upcoming(customer=cid, subscription=sid)
    except Exception:
        pass
    # Fallback: call the endpoint directly
    try:
        return stripe.Invoice._static_request("get", "/v1/invoices/upcoming", params={"customer": cid, "subscription": sid})
    except Exception:
        return None

def _stripe_subscription_raw(subscription_id: str):
    """Fetch raw subscription JSON via low-level request (works across stripe library versions)."""
    sid = (subscription_id or "").strip()
    if not sid:
        return None
    try:
        return stripe.Subscription._static_request("get", f"/v1/subscriptions/{sid}", params={})
    except Exception:
        return None

def _stripe_latest_invoice_for_subscription(subscription_id: str):
    """Fetch latest invoice for a subscription (best-effort across stripe versions)."""
    sid = (subscription_id or "").strip()
    if not sid:
        return None
    try:
        if hasattr(stripe, "Invoice") and hasattr(stripe.Invoice, "list"):
            invs = stripe.Invoice.list(subscription=sid, limit=1)
            data = list(getattr(invs, "data", []) or [])
            if data:
                return data[0]
    except Exception:
        pass
    try:
        res = stripe.Invoice._static_request("get", "/v1/invoices", params={"subscription": sid, "limit": 1})
        data = _stripe_obj_get(res, "data", []) or []
        if data:
            return data[0]
    except Exception:
        return None
    return None

def _get_stripe_price_id(plan_id: str) -> Optional[str]:
    """Map internal plan IDs to Stripe Price IDs via env vars."""
    if plan_id == 'trial_14d':
        # Trial should be a subscription (auto-converts to monthly unless canceled).
        # Use a recurring monthly price here (or a dedicated trial recurring price).
        return (
            (os.getenv('STRIPE_PRICE_TRIAL_RECURRING') or '').strip()
            or (os.getenv('STRIPE_PRICE_MONTHLY_10_95') or '').strip()
            or None
        )
    if plan_id == 'monthly_10_95':
        return (os.getenv('STRIPE_PRICE_MONTHLY_10_95') or '').strip() or None
    if plan_id == 'annual_6_95':
        return (os.getenv('STRIPE_PRICE_ANNUAL_6_95') or '').strip() or None
    return None


def _get_stripe_trial_upfront_fee_price_id() -> Optional[str]:
    """Optional one-time fee charged at checkout for the trial (e.g. $1.85).

    Create a one-time Price in Stripe and set STRIPE_PRICE_TRIAL_FEE_1_85 to its price_ id.
    """
    return (os.getenv('STRIPE_PRICE_TRIAL_FEE_1_85') or '').strip() or None


def _get_stripe_payment_link(plan_id: str) -> Optional[str]:
    """Optional Stripe Payment Links (non-secret). If set, /checkout can redirect here directly."""
    defaults = {
        # Provided by user
        'trial_14d': 'https://buy.stripe.com/cNi6oBeZ21gXfAu1cD7Vm02',
        'monthly_10_95': 'https://buy.stripe.com/bJe7sFcQU4t93RM7B17Vm03',
        'annual_6_95': 'https://buy.stripe.com/aFa9ANdUYgbR9c6bRh7Vm04',
    }
    if plan_id == 'trial_14d':
        return (os.getenv('STRIPE_PAYMENTLINK_TRIAL_14D') or '').strip() or defaults['trial_14d']
    if plan_id == 'monthly_10_95':
        return (os.getenv('STRIPE_PAYMENTLINK_MONTHLY_10_95') or '').strip() or defaults['monthly_10_95']
    if plan_id == 'annual_6_95':
        return (os.getenv('STRIPE_PAYMENTLINK_ANNUAL_6_95') or '').strip() or defaults['annual_6_95']
    return None


def _redirect_to_stripe_payment_link(plan_id: str) -> Optional['Response']:
    """Redirect to Stripe Payment Link with useful prefill params so webhook can map back to user."""
    # Trial must be a subscription with a 14-day trial and auto-convert to monthly unless canceled.
    # If Trial is a one-time Payment Link, it cannot auto-renew. So do NOT use a Payment Link for trial.
    if plan_id == 'trial_14d':
        return None
    link = _get_stripe_payment_link(plan_id)
    if not link:
        return None
    params = {}
    try:
        email = (getattr(current_user, 'email', '') or '').strip()
        if email:
            params['prefilled_email'] = email
    except Exception:
        pass
    # This is the most reliable way for webhook to map checkout back to our user.
    try:
        params['client_reference_id'] = str(getattr(current_user, 'id', '') or '')
    except Exception:
        pass
    # Helpful for debugging / analytics
    params['metadata[plan_id]'] = plan_id

    sep = '&' if ('?' in link) else '?'
    url = link + (sep + urlencode(params)) if params else link
    return redirect(url, code=303)


@app.route("/checkout")
def checkout():
    """Checkout entrypoint.

    If STRIPE_SECRET_KEY is configured, creates a Stripe Checkout Session and redirects to Stripe.
    Otherwise falls back to the placeholder confirmation page.
    """
    plan_id = request.args.get('plan', '').strip()
    plan = _get_plan_config(plan_id)
    if not plan:
        flash("Please select a valid plan.", "danger")
        return redirect(url_for("plans"))

    if not current_user.is_authenticated:
        # Require login so we can unlock paid features for the correct user.
        return redirect(url_for("login", next=request.full_path))

    # Upgrade during trial:
    # Charge the customer now (so they enter card + pay immediately), but keep the trial time.
    # We do this by charging a one-time amount now and then applying it as a customer-balance credit,
    # while updating the existing trial subscription to the target plan (annual/monthly).
    if _stripe_enabled() and plan_id in ("monthly_10_95", "annual_6_95"):
        try:
            prof = get_user_profile_azure(getattr(current_user, "id", "")) or {}
            customer_id = str(prof.get("stripe_customer_id") or "").strip()
            if not customer_id:
                customer_id = _find_stripe_customer_id_by_email((getattr(current_user, "email", "") or "").strip())
            if customer_id:
                trial_sub = _find_trialing_subscription_for_customer(customer_id)
                if trial_sub:
                    price_id = _get_stripe_price_id(plan_id)
                    if not price_id:
                        flash("Checkout is not configured. Please contact support.", "danger")
                        return redirect(url_for("plans"))
                    stripe.api_key = (os.getenv("STRIPE_SECRET_KEY") or "").strip()
                    price = stripe.Price.retrieve(price_id)
                    unit_amount = int(getattr(price, "unit_amount", 0) or 0)
                    currency = str(getattr(price, "currency", "usd") or "usd")
                    if unit_amount <= 0:
                        flash("Checkout is not configured. Please contact support.", "danger")
                        return redirect(url_for("plans"))

                    success_url = url_for('my_revisions', _external=True, _scheme=request.scheme) + "?checkout=success"
                    cancel_url = url_for('plans', _external=True, _scheme=request.scheme)
                    session_obj = stripe.checkout.Session.create(
                        mode="payment",
                        customer=customer_id,
                        line_items=[
                            {
                                "price_data": {
                                    "currency": currency,
                                    "unit_amount": unit_amount,
                                    "product_data": {"name": f"{plan.get('label')} (starts after trial)"},
                                },
                                "quantity": 1,
                            }
                        ],
                        client_reference_id=str(current_user.id),
                        metadata={
                            "upgrade_from_trial": "1",
                            "plan_id": plan_id,
                            "trial_subscription_id": str(getattr(trial_sub, "id", "") or ""),
                        },
                        success_url=success_url,
                        cancel_url=cancel_url,
                    )
                    return redirect(session_obj.url, code=303)
        except Exception:
            pass

    # Preferred: Stripe Payment Links (fastest, no API calls required here).
    pl_redirect = _redirect_to_stripe_payment_link(plan_id)
    if pl_redirect:
        return pl_redirect

    # Alternative: Real Stripe Checkout flow via API (requires secret key + price ids).
    if _stripe_enabled():
        price_id = _get_stripe_price_id(plan_id)
        if not price_id:
            flash("Checkout is not configured. Please contact support.", "danger")
            return redirect(url_for("plans"))

        stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()

        # Build absolute URLs
        success_url = url_for('my_revisions', _external=True, _scheme=request.scheme) + "?checkout=success"
        cancel_url = url_for('plans', _external=True, _scheme=request.scheme)
        try:
            subscription_data = None
            line_items = [{"price": price_id, "quantity": 1}]
            if plan_id == 'trial_14d':
                # 14-day trial that converts into the recurring monthly subscription unless canceled
                subscription_data = {"trial_period_days": 14}
                # Optional one-time upfront trial fee (e.g. $1.85)
                fee_price = _get_stripe_trial_upfront_fee_price_id()
                if fee_price:
                    line_items.append({"price": fee_price, "quantity": 1})

            session_obj = stripe.checkout.Session.create(
                mode="subscription",
                line_items=line_items,
                customer_email=(getattr(current_user, 'email', '') or None),
                client_reference_id=str(current_user.id),
                metadata={"plan_id": plan_id},
                subscription_data=subscription_data,
                success_url=success_url,
                cancel_url=cancel_url,
                allow_promotion_codes=True,
            )
            return redirect(session_obj.url, code=303)
        except Exception as e:
            logger.error(f"Stripe checkout session create failed: {str(e)}")
            flash("Checkout is temporarily unavailable. Please try again.", "danger")
            return redirect(url_for("plans"))

    # Fallback placeholder confirmation page (no Stripe configured)
    current_year = datetime.now().year
    return render_template("checkout.html", year=current_year, user=current_user, plan=plan)


@app.route("/checkout/complete", methods=["POST"])
@login_required
def checkout_complete():
    """Simulate purchase completion by updating the Azure Users profile.

    This makes plan links functional without integrating a payment processor yet.
    """
    plan_id = request.form.get('plan', '').strip()
    plan = _get_plan_config(plan_id)
    if not plan:
        flash("Invalid plan selection.", "danger")
        return redirect(url_for("plans"))

    try:
        paid_until = (datetime.now(timezone.utc) + timedelta(days=int(plan.get('duration_days') or 0))).isoformat()
        table_client = get_users_table_client()
        entity = {
            'PartitionKey': str(current_user.id),
            'RowKey': 'profile',
            'is_paid': True,
            'plan_status': plan.get('plan_status') or 'paid',
            'paid_until': paid_until,
        }
        table_client.upsert_entity(entity, mode=UpdateMode.MERGE)
        flash(f"You're all set! {plan.get('label')} activated.", "success")
        return redirect(url_for("my_revisions"))
    except Exception as e:
        logger.error(f"checkout_complete error: {str(e)}")
        flash("We couldn't activate your plan. Please try again.", "danger")
        return redirect(url_for("plans"))


@app.route("/stripe/webhook", methods=["POST"])
def stripe_webhook():
    """Stripe webhook handler.

    Required env vars:
    - STRIPE_SECRET_KEY
    - STRIPE_WEBHOOK_SECRET
    - STRIPE_PRICE_MONTHLY_10_95 / STRIPE_PRICE_ANNUAL_6_95

    Trial setup (to auto-convert to monthly unless canceled):
    - STRIPE_PRICE_TRIAL_RECURRING (optional, otherwise uses STRIPE_PRICE_MONTHLY_10_95)
    - STRIPE_PRICE_TRIAL_FEE_1_85 (optional one-time fee price)
    """
    payload = request.data
    sig_header = request.headers.get("Stripe-Signature", "")
    webhook_secret = (os.getenv("STRIPE_WEBHOOK_SECRET") or "").strip()
    if not webhook_secret:
        return ("Webhook not configured", 400)

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except Exception as e:
        logger.error(f"stripe_webhook signature error: {str(e)}")
        return ("Invalid signature", 400)

    try:
        etype = event.get("type")
        data = (event.get("data") or {}).get("object") or {}

        # We primarily rely on checkout.session.completed to map the customer to our user_id.
        if etype == "checkout.session.completed":
            client_ref = data.get("client_reference_id")
            customer_id = data.get("customer")
            subscription_id = data.get("subscription")
            plan_id = (data.get("metadata") or {}).get("plan_id", "")
            upgrade_from_trial = str((data.get("metadata") or {}).get("upgrade_from_trial") or "").strip()
            trial_subscription_id = str((data.get("metadata") or {}).get("trial_subscription_id") or "").strip()
            mode = str(data.get("mode") or "").strip().lower()

            # Fallback mapping: if client_reference_id wasn't present, try to find user by email.
            if not client_ref:
                try:
                    email = (
                        (data.get("customer_details") or {}).get("email")
                        or data.get("customer_email")
                        or ""
                    )
                    email = str(email).strip().lower()
                    if email:
                        for uid, u in users.items():
                            if (getattr(u, "email", "") or "").strip().lower() == email:
                                client_ref = uid
                                break
                except Exception:
                    pass

            if not client_ref:
                return ("No client_reference_id", 200)

            # Special case: trial upgrade where we charged up-front (payment mode) and need to:
            # 1) credit the customer balance so the first subscription invoice at trial end is covered
            # 2) update the existing trial subscription to the selected plan (annual/monthly)
            if upgrade_from_trial == "1" and mode == "payment" and customer_id and trial_subscription_id and plan_id in ("monthly_10_95", "annual_6_95"):
                stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
                try:
                    amount_total = int(data.get("amount_total") or 0)
                    currency = str(data.get("currency") or "usd")
                except Exception:
                    amount_total = 0
                    currency = "usd"

                # Apply customer balance credit equal to the amount paid now.
                # This will be applied automatically to the subscription invoice at trial end.
                try:
                    if amount_total > 0:
                        stripe.Customer.create_balance_transaction(
                            customer_id,
                            amount=-amount_total,
                            currency=currency,
                            description=f"Prepayment credit for {plan_id} (paid during trial)",
                        )
                except Exception:
                    pass

                # Update the existing trial subscription to the target plan (keeps the same trial_end).
                try:
                    price_id = _get_stripe_price_id(plan_id)
                    if price_id:
                        sub = stripe.Subscription.retrieve(trial_subscription_id, expand=["items.data"])
                        items = getattr(sub, "items", None)
                        items_data = getattr(items, "data", []) if items else []
                        item_id = str(getattr(items_data[0], "id", "") or "") if items_data else ""
                        if not item_id:
                            try:
                                si = stripe.SubscriptionItem.list(subscription=trial_subscription_id, limit=1)
                                si_data = list(getattr(si, "data", []) or [])
                                if si_data:
                                    item_id = str(getattr(si_data[0], "id", "") or "").strip()
                            except Exception:
                                item_id = ""
                        if item_id:
                            stripe.Subscription.modify(
                                trial_subscription_id,
                                items=[{"id": item_id, "price": price_id}],
                                proration_behavior="none",
                                metadata={"plan_id": plan_id},
                            )
                        # Update Azure profile (paid_until stays at trial end until renewal)
                        paid_until = ""
                        try:
                            current_period_end = getattr(sub, "current_period_end", None)
                            if current_period_end:
                                paid_until = datetime.fromtimestamp(int(current_period_end), tz=timezone.utc).isoformat()
                        except Exception:
                            paid_until = ""
                        try:
                            table_client = get_users_table_client()
                            entity = {
                                "PartitionKey": str(client_ref),
                                "RowKey": "profile",
                                "is_paid": True,
                                "plan_status": plan_id,
                                "stripe_customer_id": str(customer_id),
                                "stripe_subscription_id": str(trial_subscription_id),
                            }
                            if paid_until:
                                entity["paid_until"] = paid_until
                            table_client.upsert_entity(entity, mode=UpdateMode.MERGE)
                        except Exception:
                            pass
                except Exception:
                    pass

                return ("OK", 200)

            # Fetch subscription to compute paid_until
            stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
            paid_until = ""
            plan_status = ""
            try:
                if subscription_id and stripe.api_key:
                    sub = stripe.Subscription.retrieve(subscription_id)
                    plan_status = str(getattr(sub, "status", "") or "")
                    current_period_end = getattr(sub, "current_period_end", None)
                    if current_period_end:
                        paid_until = datetime.fromtimestamp(int(current_period_end), tz=timezone.utc).isoformat()
            except Exception:
                pass

            try:
                table_client = get_users_table_client()
                entity = {
                    "PartitionKey": str(client_ref),
                    "RowKey": "profile",
                    "is_paid": True,
                    "plan_status": (plan_id or plan_status or "paid"),
                }
                if paid_until:
                    entity["paid_until"] = paid_until
                if customer_id:
                    entity["stripe_customer_id"] = str(customer_id)
                if subscription_id:
                    entity["stripe_subscription_id"] = str(subscription_id)
                table_client.upsert_entity(entity, mode=UpdateMode.MERGE)
            except Exception as e:
                logger.error(f"stripe_webhook upsert error: {str(e)}")
                return ("Error", 500)

            # If the customer previously started a trial subscription, and then purchased a paid plan,
            # Stripe Payment Links/Checkout can result in multiple subscriptions. To avoid double-billing,
            # cancel any other *trialing* subscriptions for this customer (keep the newly purchased one).
            try:
                if customer_id and stripe.api_key:
                    subs_trialing = stripe.Subscription.list(customer=customer_id, status="trialing", limit=20)
                    tdata = list(getattr(subs_trialing, "data", []) or [])
                    for s in tdata:
                        sid = str(getattr(s, "id", "") or "")
                        if sid and subscription_id and sid == str(subscription_id):
                            continue
                        if sid:
                            try:
                                stripe.Subscription.delete(sid)
                            except Exception:
                                pass
            except Exception:
                pass

        # Keep subscription status in sync (cancel/expire)
        if etype in ("customer.subscription.updated", "customer.subscription.deleted"):
            stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
            sub = data
            customer_id = sub.get("customer")
            subscription_id = sub.get("id")
            status = sub.get("status")
            current_period_end = sub.get("current_period_end")
            paid_until = ""
            if current_period_end:
                try:
                    paid_until = datetime.fromtimestamp(int(current_period_end), tz=timezone.utc).isoformat()
                except Exception:
                    paid_until = ""

            # Find user by stripe_customer_id in Azure Users table
            if customer_id:
                try:
                    table_client = get_users_table_client()
                    # scan profiles (small scale). For large scale, add an index.
                    for e in table_client.list_entities():
                        if e.get("RowKey") != "profile":
                            continue
                        if str(e.get("stripe_customer_id") or "") == str(customer_id):
                            uid = str(e.get("PartitionKey"))
                            entity = {"PartitionKey": uid, "RowKey": "profile"}
                            entity["plan_status"] = str(status or "")
                            if paid_until:
                                entity["paid_until"] = paid_until
                            # Consider user paid only if active/trialing
                            entity["is_paid"] = str(status or "").lower() in ("active", "trialing")
                            if subscription_id:
                                entity["stripe_subscription_id"] = str(subscription_id)
                            table_client.upsert_entity(entity, mode=UpdateMode.MERGE)
                            break
                except Exception as e:
                    logger.error(f"stripe_webhook subscription sync error: {str(e)}")
                    # do not fail webhook

        return ("OK", 200)
    except Exception as e:
        logger.error(f"stripe_webhook error: {str(e)}")
        return ("Error", 500)


@app.route("/billing/portal")
@login_required
def billing_portal():
    """Send the logged-in user to Stripe Customer Portal so they can cancel/manage their plan."""
    if not _stripe_enabled():
        flash("Billing portal is not configured.", "danger")
        return redirect(url_for("plans"))

    stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
    customer_id = _get_stripe_customer_id_from_azure(getattr(current_user, 'id', ''))
    if not customer_id:
        flash("We couldn't find your billing profile yet. If you just purchased, refresh and try again.", "danger")
        return redirect(url_for("my_revisions"))

    try:
        return_url = url_for("my_revisions", _external=True, _scheme=request.scheme)
        session_obj = stripe.billing_portal.Session.create(
            customer=customer_id,
            return_url=return_url,
        )
        return redirect(session_obj.url, code=303)
    except Exception as e:
        logger.error(f"billing_portal error: {str(e)}")
        flash("Unable to open billing portal. Please try again.", "danger")
        return redirect(url_for("my_revisions"))



@app.route("/results", methods=["POST"])
#login_required
def results_route():
    print("=== results_route called ===")
    try:
        resume_text = ""
        
        # Enforce free tier revision limit (2) for authenticated non-paid users.
        if current_user.is_authenticated and (not is_paid_user(current_user)):
            try:
                used = len(get_user_revisions(current_user.id))
                if used >= FREE_REVISION_LIMIT:
                    flash("Free tier includes 2 resume revisions. Upgrade to unlock unlimited revisions and PDF downloads.", "danger")
                    return redirect(url_for("plans", limit="1"))
            except Exception:
                # If counting fails, do not block.
                pass
        
        # Check if file was uploaded
        if 'resumeFile' in request.files:
            file = request.files['resumeFile']
            if file and file.filename:
                print(f"Processing uploaded file: {file.filename}")
                try:
                    resume_text = extract_text_from_file(file)
                    print(f"Extracted text length: {len(resume_text)}")
                    if not resume_text or not resume_text.strip():
                        logger.warning("Uploaded file parsed but contained no extractable text")
                        flash("The text could not be extracted from the file possibly due to its formatting. Another option is to copy the resume text and paste it in the provided text area.", 'danger')
                        return redirect(url_for('index', scroll_to_form='true'))
                except Exception as e:
                    logger.error(f"Upload processing failed: {str(e)}")
                    flash("Your file could not be processed possibly due to its formatting. Another option is to copy  the resume text and paste it in the provided text area.", 'danger')
                    return redirect(url_for('index', scroll_to_form='true'))
        
        # If no file uploaded, check for text input  
        if not resume_text:
            resume_text = request.form.get("resume", "").strip()
        
        # Validate that we have resume content
        if not resume_text:
            flash("Please upload a resume file or paste your resume content", 'danger')
            return redirect(url_for('index', scroll_to_form='true'))
        
        # Get job description (optional)
        job_description = request.form.get("jobDescription", "").strip()
        
        # Process the resume
        print(f"Resume text preview (first 200 chars): {resume_text[:200]}...")
        print(f"Job description preview: {job_description[:100] if job_description else 'None'}...")
        revised_resume, feedback = revise_resume(resume_text, job_description)
        print("=== FEEDBACK DATA ===")
        print(f"Feedback type: {type(feedback)}")
        print(json.dumps(feedback, indent=2))
        
        # Save to user account if authenticated
        if current_user.is_authenticated:
            import uuid
            try:
                save_resume_revision(
                    current_user.id,
                    str(uuid.uuid4()),
                    revised_resume,
                    feedback=feedback,
                    original_resume=resume_text,
                    job_description=job_description
                )
            except FreeTierLimitReached:
                # Still show results, but do not persist a new revision.
                flash("You've reached the free tier limit (2 revisions). Upgrade to save unlimited revisions.", "danger")
            except Exception:
                pass
        else:
            # Store revision in session for post-signup saving
            session['pending_revision'] = {
                'revised_resume': revised_resume,
                'feedback': feedback,
                'original_resume': resume_text,
                'job_description': job_description
            }
        
        # Track conversion (resume submission)
        conversion_info = analytics.track_conversion(session, "resume_submission")
        print(f"=== CONVERSION TRACKED ===")
        print(f"Conversion info: {conversion_info}")
        
        # Store data in session and redirect (Post/Redirect/Get) to prevent resubmission on back
        session['results_data'] = {
            'original_resume': resume_text,
            'revised_resume': revised_resume,
            'feedback': feedback,
            'job_description': job_description,
        }
        # Ensure session persistence + clear stale template data (prevents old template snapshot confusion)
        session.pop('template_data', None)
        session.modified = True
        # Bust caches (browser/service worker) for /results
        import time
        return redirect(url_for('results_get', ts=int(time.time())))
    except Exception as e:
        import traceback
        print("=== ERROR IN revise_resume_route ===")
        print(f"Error type: {type(e).__name__}")
        print(f"Error message: {str(e)}")
        traceback.print_exc()
        print("=== END ERROR DEBUG ===")
        flash(f"Error: {str(e)}", 'danger')
        return redirect(url_for('index'))




@app.route("/results", methods=["GET"])
def results_get():
    data = session.get('results_data', None)
    if not data:
        # No data to show; notify and send the user to the homepage
        flash("⚠️ Your file could not be processed possibly due to its formatting. Please attempt to copy the resume text and paste it in the provided text area.", 'danger')
        return redirect(url_for('index'))
    # Keep data in session for template selection
    # session.pop would remove it, so we use session.get and keep it available
    resp = make_response(render_template(
        "result.html",
        original_resume=data.get('original_resume', ''),
        revised_resume=data.get('revised_resume', ''),
        feedback=data.get('feedback', {}),
        job_description=data.get('job_description', ''),
        error=None
    ))
    # Transactional page: never cache (prevents showing a previous resume after a new submission)
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    resp.headers['Vary'] = 'Cookie'
    return resp

@app.route("/api/parse-resume-for-template", methods=["POST"])
def parse_resume_for_template():
    """Parse resume and store structured data in session for template viewing"""
    try:
        data = request.get_json()
        template_name = data.get('template', 'modern')
        
        # Validate template name
        valid_templates = [
            'modern',
            'professional',
            'minimal',
            'creative',
            'classicSidebar',
            'classic-sidebar',
            'classic_sidebar',
            'classicPortrait',
            'classic-portrait',
            'classic_portrait',
            'darkSidebar',
            'dark-sidebar',
            'dark_sidebar',
            'darkSidebarProgress',
            'dark-sidebar-progress',
            'dark_sidebar_progress',
            'navyHeader',
            'navy-header',
            'navy_header',
            'timelineBlue',
            'timeline-blue',
            'timeline_blue',
            'oliveClassic',
            'olive-classic',
            'olive_classic',
        ]
        if template_name not in valid_templates:
            return jsonify({"success": False, "error": "Invalid template name"}), 400
        
        # Get revised resume from session
        results_data = session.get('results_data')
        if not results_data:
            return jsonify({"success": False, "error": "Resume data not found"}), 404
        
        revised_resume = results_data.get('revised_resume', '')
        if not revised_resume:
            return jsonify({"success": False, "error": "Revised resume not found"}), 404
        
        # Parse resume to get structured data
        parsed_result = parse_resume(revised_resume)
        structured_resume = parsed_result.get('resume', {})
        
        # Store in session for template viewer
        session['template_data'] = {
            'structured_resume': structured_resume,
            'template_name': template_name,
            'revised_resume': revised_resume  # Keep original text as fallback
        }
        
        return jsonify({
            "success": True,
            "template": template_name
        })
        
    except Exception as e:
        logger.error(f"Error parsing resume for template: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/template-data", methods=["GET"])
def get_template_data():
    """Get structured resume data for template viewer"""
    template_data = session.get('template_data')
    if not template_data:
        return jsonify({"error": "Template data not found"}), 404
    
    return jsonify({
        "success": True,
        "resume": template_data['structured_resume'],
        "template": template_data['template_name'],
        "revised_resume": template_data.get('revised_resume', '')
    })


@app.route("/api/template-data", methods=["POST"])
def update_template_data():
    """Update structured resume data for template viewer (stored in session)."""
    try:
        template_data = session.get('template_data')
        if not template_data:
            return jsonify({"success": False, "error": "Template data not found"}), 404

        data = request.get_json(force=True, silent=True) or {}
        resume = data.get("resume")
        if not isinstance(resume, dict):
            return jsonify({"success": False, "error": "Invalid resume payload"}), 400

        # Update only the structured resume. Keep template_name and revised_resume intact.
        template_data["structured_resume"] = resume
        session["template_data"] = template_data
        session.modified = True

        return jsonify({"success": True})
    except Exception as e:
        logger.error(f"Error updating template data: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/termsprivacy")
def termsprivacy():
    return redirect(url_for('terms_privacy'), code=301)

@app.route("/terms_privacy")
def terms_privacy():
    return render_template("terms_privacy.html")

@app.route("/privacy")
def privacy():
    return render_template("privacy.html")



@app.route("/about")
def about():
    current_year = datetime.now().year
    return render_template("about.html", year=current_year, user=current_user if current_user.is_authenticated else None)

@app.route("/blog")
def blog():
    current_year = datetime.now().year
    return render_template("blog.html", year=current_year, user=current_user if current_user.is_authenticated else None)

# Blog post metadata for SEO optimization
BLOG_POSTS_METADATA = {
    'ats-optimization': {
        'title': 'How to Beat the ATS in 2025 | ResumaticAI',
        'meta_description': 'Master the art of creating ATS-friendly resumes with our comprehensive guide. Learn how Applicant Tracking Systems work and discover proven strategies to ensure your resume gets past automated screening.',
        'keywords': 'ATS optimization, applicant tracking system, resume keywords, ATS friendly resume, job application tips, resume writing',
        'og_title': 'How to Beat the ATS in 2025 - Complete Guide',
        'og_description': 'Learn proven strategies to create ATS-friendly resumes that pass automated screening systems and get you more job interviews.',
        'og_image': 'https://resumaticai.com/static/images/robot.png'
    },
    'Resume-objective': {
        'title': 'Resume Summary vs. Objective: Which Should You Use in 2025? | ResumaticAI',
        'meta_description': 'Learn the key differences between resume summaries and objectives, and discover which one will help you land more interviews in 2025.',
        'keywords': 'resume summary, resume objective, resume writing tips, career advice, job application, resume format',
        'og_title': 'Resume Summary vs. Objective: 2025 Guide',
        'og_description': 'Discover whether to use a resume summary or objective statement in 2025 to maximize your interview chances.',
        'og_image': 'https://resumaticai.com/static/images/vs.png'
    },
    'Power-words': {
        'title': 'Power Words to Use in Your Resume (And Ones to Avoid) | ResumaticAI',
        'meta_description': 'Discover the most impactful words to use in your resume and learn which overused phrases to avoid to make your application stand out.',
        'keywords': 'resume power words, action verbs, resume writing, job application, career tips, resume optimization',
        'og_title': 'Resume Power Words: What to Use and Avoid',
        'og_description': 'Learn which words will make your resume stand out and which ones to avoid for better job application success.',
        'og_image': 'https://resumaticai.com/static/images/power.jpg'
    },
    'no-experience': {
        'title': 'How to Write a Resume With No Work Experience | ResumaticAI',
        'meta_description': 'Get expert tips on creating a compelling resume when you\'re just starting out, with strategies to highlight your skills and potential.',
        'keywords': 'resume no experience, first resume, entry level resume, student resume, career starter, resume writing',
        'og_title': 'Resume Writing for Beginners: No Experience Needed',
        'og_description': 'Create a compelling resume even without work experience using our expert strategies and tips.',
        'og_image': 'https://resumaticai.com/static/images/noexperience.jpg'
    },
    'resume-format-2025': {
        'title': 'Best Resume Formats for 2025 (With Examples) | ResumaticAI',
        'meta_description': 'Explore the most effective resume formats for 2025, complete with real examples and guidelines for different career stages.',
        'keywords': 'resume format 2025, resume templates, resume examples, resume layout, career stages, resume design',
        'og_title': 'Best Resume Formats for 2025: Complete Guide',
        'og_description': 'Choose the perfect resume format for 2025 with our comprehensive guide and real examples.',
        'og_image': 'https://resumaticai.com/static/images/format.webp'
    },
    'resume-format-2026': {
        'title': 'Best Resume Formats for 2026 (With Examples) | ResumaticAI',
        'meta_description': 'See the best resume formats for 2026 with ATS-friendly examples. Learn when to use reverse-chronological vs. combination formats, how to structure skills, and whether to use PDF or DOCX.',
        'keywords': 'resume format 2026, resume templates, resume examples, resume layout, career stages, resume design',
        'og_title': 'Best Resume Formats for 2026: Complete Guide',
        'og_description': 'ATS-friendly resume formats, layout rules, and examples for 2026.',
        'og_image': 'https://resumaticai.com/static/images/format.webp'
    },
    'toptenmistakes': {
        'title': 'Top 10 Resume Mistakes to Avoid in 2025 | ResumaticAI',
        'meta_description': 'Learn the most common resume mistakes that can cost you job opportunities and how to avoid them in 2025.',
        'keywords': 'resume mistakes, resume errors, job application tips, resume writing, career advice, avoid resume mistakes',
        'og_title': 'Top 10 Resume Mistakes to Avoid in 2025',
        'og_description': 'Don\'t let these common resume mistakes cost you job opportunities. Learn how to avoid them.',
        'og_image': 'https://resumaticai.com/static/images/robot.png'
    },
    'tailorresumejob': {
        'title': 'How to Tailor Your Resume for Each Job Application | ResumaticAI',
        'meta_description': 'Learn proven strategies to customize your resume for each job application to increase your chances of getting interviews and job offers.',
        'keywords': 'tailor resume, customize resume, job application, resume customization, targeted resume, job-specific resume',
        'og_title': 'How to Tailor Your Resume for Each Job Application',
        'og_description': 'Master the art of customizing your resume for each job to maximize your interview chances and career success.',
        'og_image': 'https://resumaticai.com/static/images/robot.png'
    }
}

@app.route("/blog/<post>")
def blog_post(post):
    current_year = datetime.now().year

    # Canonicalize blog slugs to prevent 500s from template mismatches
    # Accept underscores/case-insensitive inputs and redirect to canonical slugs
    canonical_slug_map = {
        # key: normalized (lowercase, hyphens) -> value: canonical file/slug
        "toptenmistakes": "toptenmistakes",
        "ats-optimization": "ats-optimization",
        "resume-objective": "Resume-objective",
        "resume-summary-examples": "Resume-objective",
        "no-experience": "no-experience",
        "power-words": "Power-words",
    "resume-format-2025": "resume-format-2026",
    "resume-format-2026": "resume-format-2026",
        "tailorresumejob": "tailorresumejob",
    }

    requested_slug = post.strip()
    normalized_slug = requested_slug.replace("_", "-").lower()

    if normalized_slug in canonical_slug_map:
        canonical_slug = canonical_slug_map[normalized_slug]
        # Redirect variants to the canonical version for SEO consistency
        if requested_slug != canonical_slug:
            return redirect(url_for("blog_post", post=canonical_slug), code=301)
    else:
        # Unknown slug -> 404 instead of template error 500
        from flask import abort
        abort(404)

    # Get metadata for the blog post using the canonical slug
    post_metadata = BLOG_POSTS_METADATA.get(canonical_slug, {})

    try:
        return render_template(
            f"{canonical_slug}.html",
            year=current_year,
            user=current_user if current_user.is_authenticated else None,
            post_metadata=post_metadata,
        )
    except TemplateNotFound:
        # Fallback: if 2026 slug is requested but template file not present, use 2025 template
        file_slug = canonical_slug
        if canonical_slug == 'resume-format-2026':
            file_slug = 'resume-format-2025'
        return render_template(
            f"{file_slug}.html",
            year=current_year,
            user=current_user if current_user.is_authenticated else None,
            post_metadata=post_metadata,
        )

@app.route("/templates")
def templates():
    # Explicitly pass API key to handle Azure environment
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        raise ValueError("OPENAI_API_KEY environment variable not set")
    
    # Configure timeout and retry for Azure reliability
    client_templates = OpenAI(
        api_key=api_key,
        timeout=60.0,
        max_retries=3
    )
    prompt= f'''Please create a python dictionary with the section names of the provided resume 
    as keys and the associated content as values.'''
    try:
            response = client_templates.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}]
            )
    except Exception as e:
            error_msg = str(e)
            print(f"OpenAI API Error: {error_msg}")
            logger.error(f"OpenAI API Error details: {type(e).__name__}: {error_msg}")
            raise ValueError(f"Failed to connect to OpenAI API: {error_msg}")

    resume_data = response.choices[0].message.content
    resume_data = request.json.get('resume_data')

    return render_template('resume_template.html', resume=resume_data)


@app.route("/subscribe", methods=["POST"])
def subscribe():
    email = request.form.get("email")
    referrer = request.referrer or url_for("index")
    
    if email:
        try:
            # Create subscribers.csv if it doesn't exist
            import os
            if not os.path.exists("subscribers.csv"):
                with open("subscribers.csv", "w") as f:
                    f.write("email\n")  # Header row
            
            # Check if email already exists
            existing_emails = []
            try:
                with open("subscribers.csv", "r") as f:
                    existing_emails = [line.strip().lower() for line in f.readlines()]
            except FileNotFoundError:
                pass
            
            if email.lower() in existing_emails:
                flash("You're already subscribed to our newsletter!", "info")
            else:
                with open("subscribers.csv", "a") as f:
                    f.write(email + "\n")
                flash("Thank you for subscribing! You'll receive our latest resume tips and career advice.", "success")
            
            # Smart redirect based on referrer
            if "/blog" in referrer:
                return redirect(url_for("blog"))
            elif "/thank_you" in referrer:
                return redirect(url_for("thank_you"))
            else:
                return redirect(url_for("thank_you"))
        except Exception as e:
            flash("Something went wrong. Please try again later.", "danger")
            print(f"Subscription error: {str(e)}")
            return redirect(referrer)
    else:
        flash("Please enter a valid email address.", "danger")
        return redirect(referrer)

@app.route("/thank_you")
def thank_you():
    return render_template("thank_you.html")

@app.route("/users")
@login_required
def view_users():
    if not current_user.is_authenticated:
        flash("You need to log in to view this page.", "danger")
        return redirect(url_for("login"))

    # Check if the user is an admin
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to view this page.", "danger")
        return redirect(url_for("index"))

    return render_template("users.html", users=users.values())

@app.route("/analytics")
@login_required
def view_analytics():
    """Admin dashboard for viewing Facebook ad performance and site analytics."""
    if not current_user.is_authenticated:
        flash("You need to log in to view this page.", "danger")
        return redirect(url_for("login"))

    # Check if the user is an admin
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to view this page.", "danger")
        return redirect(url_for("index"))

    # Get comprehensive analytics data
    analytics_data = analytics.get_full_analytics()
    
    return render_template("analytics.html", analytics=analytics_data)

@app.route("/api/track", methods=["POST"])
def track_conversion():
    """API endpoint for tracking conversions and events from Facebook ads."""
    try:
        # Check if this is a duplicate tracking call in the same session
        if 'visit_tracked' in session:
            # Just track the conversion/event without counting as a new visit
            source_info = session.get('traffic_source', {'type': 'unknown'})
            app.logger.info("Facebook tracking call - visit already tracked, skipping duplicate")
        else:
            # Track the visit/conversion (first time)
            source_info = analytics.track_visit(request)
            session['visit_tracked'] = True
            session['traffic_source'] = source_info
            app.logger.info(f"Facebook tracking - new visit tracked: {source_info.get('type', 'unknown')}")
        
        # Return JSON response for AJAX calls
        return {
            "status": "success",
            "source_type": source_info.get("type", "unknown"),
            "campaign": source_info.get("campaign", "unknown"),
            "message": "Event tracked successfully"
        }, 200
    except Exception as e:
        return {
            "status": "error",
            "message": str(e)
        }, 500

@app.route("/api/track_visit", methods=["POST"])
def track_visit():
    """API endpoint for tracking legitimate page visits (bot-filtered)."""
    try:
        # Check if visit already tracked in this session to prevent double counting
        if 'visit_tracked' in session:
            app.logger.info("Visit already tracked in session, skipping duplicate API tracking")
            return {
                "status": "success",
                "message": "Visit already tracked in session",
                "duplicate_prevented": True
            }, 200
        
        # Get JSON data from request
        data = request.get_json()
        
        if not data:
            return {"status": "error", "message": "No data provided"}, 400
        
        # Additional server-side bot detection
        user_agent = request.headers.get('User-Agent', '').lower()
        bot_indicators = [
            'bot', 'crawl', 'spider', 'scraper', 'curl', 'wget', 'python',
            'java', 'php', 'ruby', 'go-http', 'apache', 'nginx'
        ]
        
        # Check if request looks automated
        if any(indicator in user_agent for indicator in bot_indicators):
            return {"status": "ignored", "message": "Bot detected"}, 200
        
        # Check for required browser headers that bots often miss
        if not request.headers.get('Accept-Language'):
            return {"status": "ignored", "message": "Missing browser headers"}, 200
        
        # Track the legitimate visit using existing analytics
        source_info = analytics.track_visit(request)
        session['visit_tracked'] = True
        session['traffic_source'] = source_info
        
        # Log the visit for debugging
        app.logger.info(f"API visit tracked: {data.get('page', 'unknown')} from {request.remote_addr} - {source_info.get('type', 'unknown')}")
        
        return {
            "status": "success",
            "source_type": source_info.get("type", "unknown"),
            "message": "Visit tracked successfully"
        }, 200
        
    except Exception as e:
        logger.error(f"Error tracking visit: {str(e)}")
        return {
            "status": "error",
            "message": "Internal server error"
        }, 500

@app.route("/api/visit_count", methods=["GET"])
def get_visit_count():
    """API endpoint to retrieve current visit count."""
    try:
        # Get analytics data
        analytics_data = analytics.get_full_analytics()
        
        # Extract summary data
        summary = analytics_data.get('summary', {})
        
        return {
            "status": "success",
            "data": {
                "total_visits": summary.get('total_visits', 0),
                "facebook_ad_visits": summary.get('facebook_ad_visits', 0),
                "organic_visits": summary.get('organic_visits', 0),
                "total_conversions": summary.get('total_conversions', 0),
                "facebook_ad_conversions": summary.get('facebook_ad_conversions', 0),
                "last_updated": summary.get('last_updated', 'unknown')
            }
        }, 200
        
    except Exception as e:
        logger.error(f"Error retrieving visit count: {str(e)}")
        return {
            "status": "error", 
            "message": "Failed to retrieve visit count"
        }, 500

@app.route("/admin/stats")
@login_required
def admin_stats():
    """Admin page to view visit statistics."""
    if not current_user.is_authenticated:
        flash("You need to log in to view this page.", "danger")
        return redirect(url_for("login"))

    # Check if the user is an admin
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to view this page.", "danger")
        return redirect(url_for("index"))

    try:
        # Get comprehensive analytics data
        analytics_data = analytics.get_full_analytics()
        # Load recent feedback submissions from CSV (if present)
        feedback_rows = []
        try:
            import csv
            feedback_path = 'download_feedback.csv'
            if os.path.exists(feedback_path):
                with open(feedback_path, 'r', encoding='utf-8', newline='') as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        feedback_rows.append({
                            'timestamp_iso': row.get('timestamp_iso', ''),
                            'user_id': row.get('user_id', ''),
                            'user_email': row.get('user_email', ''),
                            'rating': row.get('rating', ''),
                            'comment': row.get('comment', ''),
                            'comparison': row.get('comparison', ''),
                        })
                # Keep only the last 100, newest first
                feedback_rows = feedback_rows[-100:][::-1]
        except Exception as e:
            app.logger.warning(f"Failed to read download_feedback.csv: {str(e)}")
        
        # Debug logging to help troubleshoot
        app.logger.info(f"Analytics data structure: {type(analytics_data)}")
        if isinstance(analytics_data, dict) and 'summary' in analytics_data:
            app.logger.info(f"Summary data: {analytics_data['summary']}")
        else:
            app.logger.warning(f"Unexpected analytics data structure: {analytics_data}")
        
        return render_template("admin_stats.html", 
                             analytics=analytics_data, 
                             user=current_user,
                             feedback_rows=feedback_rows)
    except Exception as e:
        app.logger.error(f"Error loading statistics: {str(e)}")
        flash(f"Error loading statistics: {str(e)}", "danger")
        return redirect(url_for("index"))

@app.route('/admin/feedback.csv')
@login_required
def admin_feedback_csv():
    """Download raw feedback CSV (admin only)."""
    if not getattr(current_user, "is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    try:
        path = 'download_feedback.csv'
        if not os.path.exists(path):
            # Return an empty CSV with header for convenience
            from io import StringIO, BytesIO
            buf = StringIO("timestamp_iso,user_id,user_email,rating,comment,comparison\n")
            data = buf.getvalue().encode('utf-8')
            bio = BytesIO(data)
            bio.seek(0)
            return send_file(bio, mimetype='text/csv', as_attachment=True, download_name='download_feedback.csv')
        return send_file(path, mimetype='text/csv', as_attachment=True, download_name='download_feedback.csv')
    except Exception as e:
        app.logger.error(f"Failed to serve feedback CSV: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500


# Health check endpoint for Azure App Service
@app.route('/path/health', methods=['GET'])
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint for Azure App Service monitoring"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'resumatic'
    }), 200


from flask import Response 
 



@app.route('/sitemap.xml', methods=['GET'])
def sitemap():
    """Generate dynamic sitemap with only public pages"""
    pages = []
    lastmod = datetime.now().date().isoformat()

    # Define allowed public endpoints
    public_endpoints = {
        'index': {'priority': '1.0', 'changefreq': 'daily'},
        'about': {'priority': '0.8', 'changefreq': 'monthly'},
        'blog': {'priority': '0.9', 'changefreq': 'weekly'},
        'contact': {'priority': '0.7', 'changefreq': 'monthly'},
        'privacy': {'priority': '0.3', 'changefreq': 'yearly'},
        'terms_privacy': {'priority': '0.3', 'changefreq': 'yearly'},
    }

    # Add static pages
    for endpoint, config in public_endpoints.items():
        try:
            url = url_for(endpoint, _external=True, _scheme='https')
            pages.append(f"""
                <url>
                    <loc>{url}</loc>
                    <lastmod>{lastmod}</lastmod>
                    <changefreq>{config['changefreq']}</changefreq>
                    <priority>{config['priority']}</priority>
                </url>""")
        except Exception:
            # Skip if endpoint doesn't exist
            continue

    # Add blog posts (assuming they follow the pattern /blog/<post>)
    blog_posts = [
        'toptenmistakes',
        'ats-optimization', 
        'Resume-objective',
        'no-experience',
        'Power-words',
        'resume-format-2026',
        'tailorresumejob'
    ]
    
    for post in blog_posts:
        try:
            url = url_for('blog_post', post=post, _external=True, _scheme='https')
            pages.append(f"""
                <url>
                    <loc>{url}</loc>
                    <lastmod>{lastmod}</lastmod>
                    <changefreq>monthly</changefreq>
                    <priority>0.8</priority>
                </url>""")
        except Exception:
            continue

    sitemap_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
    <urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
        {''.join(pages)}
    </urlset>"""

    response = Response(sitemap_xml, mimetype='application/xml')
    response.headers['Cache-Control'] = 'public, max-age=86400'  # Cache for 24 hours
    return response


@app.route('/api/feedback/download', methods=['POST'])
@login_required
def feedback_download_api():
    """Collect quick feedback after a logged-in user clicks to download a resume."""
    try:
        data = request.get_json(silent=True) or {}
        rating = str(data.get('rating', '')).strip()
        comment = str(data.get('comment', '')).strip()
        comparison = str(data.get('comparison', '')).strip()
        # Persist minimally to CSV; avoids DB schema work
        import csv
        from datetime import datetime, timezone
        filename = 'download_feedback.csv'
        file_exists = os.path.exists(filename)
        with open(filename, 'a', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if not file_exists:
                writer.writerow(['timestamp_iso', 'user_id', 'user_email', 'rating', 'comment', 'comparison'])
            writer.writerow([
                datetime.now(timezone.utc).isoformat(),
                getattr(current_user, 'id', ''),
                getattr(current_user, 'email', ''),
                rating,
                comment[:500],
                comparison
            ])
        return jsonify({'status': 'ok'}), 200
    except Exception as e:
        logger.error(f"feedback_download_api error: {str(e)}")
        return jsonify({'status': 'error'}), 500

# Azure Table Storage setup
AZURE_TABLE_NAME = os.getenv('AZURE_TABLE_NAME', 'ResumeRevisions')
AZURE_STORAGE_ACCOUNT = os.getenv('AZURE_STORAGE_ACCOUNT')

def get_table_client(table_name: str = None):
    connection_string = os.getenv('AZURE_STORAGE_CONNECTION_STRING')
    if connection_string:
        service = TableServiceClient.from_connection_string(conn_str=connection_string)
    else:
        credential = DefaultAzureCredential()
        service = TableServiceClient(endpoint=f"https://{AZURE_STORAGE_ACCOUNT}.table.core.windows.net", credential=credential)
    table_client = service.get_table_client(table_name or AZURE_TABLE_NAME)
    try:
        table_client.create_table()
    except Exception:
        pass  # Table may already exist
    return table_client

# Azure Users table helpers
AZURE_USERS_TABLE = os.getenv('AZURE_USERS_TABLE', 'Users')

def get_users_table_client():
    return get_table_client(AZURE_USERS_TABLE)

FREE_REVISION_LIMIT = int(os.getenv('FREE_REVISION_LIMIT', '2'))
PAID_EMAILS = set([e.strip().lower() for e in (os.getenv('PAID_EMAILS', '') or '').split(',') if e.strip()])

class FreeTierLimitReached(Exception):
    pass

def _parse_iso_dt(s: str):
    if not s:
        return None
    try:
        dt = datetime.fromisoformat(s)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt
    except Exception:
        return None

def get_user_profile_azure(user_id: str) -> Optional[dict]:
    try:
        table_client = get_users_table_client()
        return table_client.get_entity(partition_key=str(user_id), row_key='profile')
    except Exception:
        return None


def _format_paid_until(paid_until_str: str) -> str:
    """Best-effort formatting for paid_until ISO string."""
    s = (paid_until_str or "").strip()
    if not s:
        return ""
    try:
        dt = datetime.fromisoformat(s.replace("Z", "+00:00"))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        # Show local-ish friendly date; keep time out for simplicity
        return dt.astimezone(timezone.utc).strftime("%b %d, %Y")
    except Exception:
        return s

def is_paid_user(user_obj: Optional['User']) -> bool:
    try:
        if not user_obj or not getattr(user_obj, 'is_authenticated', False):
            return False
        if bool(getattr(user_obj, 'is_admin', False)):
            return True
        email = (getattr(user_obj, 'email', '') or '').strip().lower()
        if email and email in PAID_EMAILS:
            return True
        prof = get_user_profile_azure(getattr(user_obj, 'id', ''))
        if not prof:
            return False
        # Supported fields (manual or future Stripe webhook):
        # - is_paid: boolean
        # - plan_status: free|trial|paid|active|canceled
        # - paid_until: ISO timestamp (optional)
        if bool(prof.get('is_paid', False)):
            return True
        plan_status = str(prof.get('plan_status') or '').strip().lower()
        if plan_status in ('paid', 'active', 'trial', 'monthly', 'annual'):
            paid_until = _parse_iso_dt(str(prof.get('paid_until') or '').strip())
            if paid_until is None:
                return True
            return paid_until >= datetime.now(timezone.utc)
        return False
    except Exception:
        return False

def is_paid_user_id(user_id: str) -> bool:
    try:
        u = users.get(str(user_id))
        if u:
            return is_paid_user(u)
        prof = get_user_profile_azure(str(user_id))
        if not prof:
            return False
        if bool(prof.get('is_paid', False)):
            return True
        plan_status = str(prof.get('plan_status') or '').strip().lower()
        if plan_status in ('paid', 'active', 'trial', 'monthly', 'annual'):
            paid_until = _parse_iso_dt(str(prof.get('paid_until') or '').strip())
            if paid_until is None:
                return True
            return paid_until >= datetime.now(timezone.utc)
        email = str(prof.get('email') or '').strip().lower()
        return bool(email and email in PAID_EMAILS)
    except Exception:
        return False

def upsert_user_profile_azure(user_obj: 'User') -> None:
    try:
        table_client = get_users_table_client()
        provider = "google" if str(user_obj.id).isdigit() else ("facebook" if str(user_obj.id).startswith("facebook_") else "other")
        entity = {
            'PartitionKey': str(user_obj.id),
            'RowKey': 'profile',
            'name': user_obj.name or '',
            'email': user_obj.email or '',
            'created_at': user_obj.created_at,
            'is_admin': bool(getattr(user_obj, 'is_admin', False)),
            'provider': provider,
        }
        table_client.upsert_entity(entity)
    except Exception as e:
        logger.error(f"Failed to upsert user profile to Azure: {str(e)}")

# Save a revision to Azure Table Storage
def save_resume_revision(user_id, revision_id, resume_content, feedback=None, original_resume=None, notes=None, job_description=None):
    from datetime import datetime, timezone
    # Enforce free tier cap (2 revisions) for non-paid users.
    # Note: We enforce here as a safety net; primary gating happens earlier in /results.
    if not is_paid_user_id(user_id):
        try:
            existing = get_user_revisions(user_id)
            if len(existing) >= FREE_REVISION_LIMIT:
                raise FreeTierLimitReached("Free tier revision limit reached")
        except FreeTierLimitReached:
            raise
        except Exception:
            # If counting fails, do not block saves.
            pass
    table_client = get_table_client()
    entity = {
        'PartitionKey': user_id,
        'RowKey': revision_id,
        'timestamp': datetime.now(timezone.utc).isoformat(),
        'resume_content': resume_content,
        'feedback': json.dumps(feedback) if feedback else '',
        'original_resume': original_resume or '',
        'notes': notes or '',
        'job_description': job_description or '',
        # Structured application tracking (list of {company, role, date_applied, status, link})
        'applications': '[]',
    }
    table_client.upsert_entity(entity)

def _parse_applications(raw):
    """Parse stored applications JSON safely into a list of dicts."""
    if not raw:
        return []
    if isinstance(raw, list):
        apps = raw
    else:
        s = str(raw)
        try:
            apps = json.loads(s) if s.strip() else []
        except Exception:
            apps = []
    if not isinstance(apps, list):
        return []
    cleaned = []
    for item in apps:
        if not isinstance(item, dict):
            continue
        cleaned.append({
            'company': str(item.get('company', '') or '')[:120],
            'role': str(item.get('role', '') or '')[:120],
            'date_applied': str(item.get('date_applied', '') or '')[:32],
            'status': str(item.get('status', '') or '')[:40],
            'link': str(item.get('link', '') or '')[:500],
        })
    return cleaned

# Get all revisions for a user
def get_user_revisions(user_id):
    table_client = get_table_client()
    entities = table_client.query_entities(f"PartitionKey eq '{user_id}'")
    revisions = []
    for e in entities:
        timestamp_str = e.get('timestamp')
        if not timestamp_str:
            timestamp_str = str(e.get('Timestamp')) if e.get('Timestamp') else None
        try:
            timestamp = datetime.fromisoformat(timestamp_str) if timestamp_str else None
            if timestamp and timestamp.tzinfo is None:
                timestamp = timestamp.replace(tzinfo=timezone.utc)
        except Exception:
            timestamp = None
        feedback = {}
        if e.get('feedback'):
            try:
                feedback = json.loads(e['feedback'])
            except Exception:
                feedback = {}
        revisions.append({
            'revision_id': e['RowKey'],
            'timestamp': timestamp,
            'resume_content': e.get('resume_content', ''),
            'feedback': feedback,
            'original_resume': e.get('original_resume', ''),
            'notes': e.get('notes', ''),
            'job_description': e.get('job_description', ''),
            'applications': _parse_applications(e.get('applications', '')),
        })
    utc_min = datetime.min.replace(tzinfo=timezone.utc)
    revisions.sort(key=lambda x: x['timestamp'] or utc_min, reverse=True)
    for r in revisions:
        apps = r.get('applications') or []
        has_apps = bool(apps)
        has_notes = bool((r.get('notes') or '').strip())
        r['has_applications'] = bool(has_apps or has_notes)
        r['applications_count'] = len(apps)
    return revisions

# Route: My Revisions
@app.route('/my_revisions')
@login_required
def my_revisions():
    revisions = get_user_revisions(current_user.id)
    return render_template('my_revisions.html', revisions=revisions, user=current_user, is_paid=is_paid_user(current_user))


@app.route('/settings')
@login_required
def settings_page():
    prof = get_user_profile_azure(getattr(current_user, 'id', '')) or {}
    paid_until_raw = str(prof.get('paid_until') or '').strip()
    customer_id = str(prof.get('stripe_customer_id') or '').strip()
    subscription_id = str(prof.get('stripe_subscription_id') or '').strip()
    plan_status_raw = str(prof.get('plan_status') or '').strip()
    debug = str(request.args.get('debug') or '').strip() == '1'
    debug_info = None

    paid_flag = bool(is_paid_user(current_user))

    # If webhook hasn't populated Stripe ids yet (or paid flag is stale), recover them via email lookup.
    # This allows newly-purchased users to see accurate plan dates immediately.
    if _stripe_enabled():
        try:
            if not customer_id:
                customer_id = _find_stripe_customer_id_by_email((getattr(current_user, 'email', '') or '').strip())
            if customer_id and not subscription_id:
                stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
                subs = stripe.Subscription.list(customer=customer_id, status='all', limit=10)
                sdata = list(getattr(subs, 'data', []) or [])
                # Prefer active > trialing > others; then later period end
                def _rank(sub):
                    status = str(getattr(sub, 'status', '') or '').lower()
                    cpe = int(getattr(sub, 'current_period_end', 0) or 0)
                    sr = 0
                    if status == 'active':
                        sr = 3
                    elif status == 'trialing':
                        sr = 2
                    elif status in ('past_due', 'unpaid'):
                        sr = 1
                    return (sr, cpe)
                if sdata:
                    best = sorted(sdata, key=_rank, reverse=True)[0]
                    subscription_id = str(getattr(best, 'id', '') or '').strip() or subscription_id
                    best_status = str(getattr(best, 'status', '') or '').strip().lower()
                    if best_status in ('active', 'trialing'):
                        paid_flag = True

            # Backfill ids to Azure so future loads are fast/stable.
            if customer_id or subscription_id:
                try:
                    table_client = get_users_table_client()
                    entity = {"PartitionKey": str(current_user.id), "RowKey": "profile"}
                    if customer_id:
                        entity["stripe_customer_id"] = str(customer_id)
                    if subscription_id:
                        entity["stripe_subscription_id"] = str(subscription_id)
                    # If Stripe indicates an active/trialing sub, reflect that.
                    if paid_flag:
                        entity["is_paid"] = True
                    table_client.upsert_entity(entity, mode=UpdateMode.MERGE)
                except Exception:
                    pass
        except Exception:
            pass

    # Stripe-derived dates (handles upgrades where Azure still reflects the trial subscription)
    stripe_dates = _get_stripe_plan_dates_for_customer(customer_id) if paid_flag else {}
    if paid_flag and not stripe_dates and subscription_id:
        stripe_dates = _get_stripe_plan_dates_for_subscription(subscription_id)
    next_billing_iso = str(stripe_dates.get('next_billing_iso') or '').strip()
    paid_through_est_iso = str(stripe_dates.get('paid_through_est_iso') or '').strip()
    interval_label = str(stripe_dates.get('interval_label') or '').strip()

    # Strongest source of truth: compute directly from the stored Stripe subscription_id.
    # This avoids cases where precomputed stripe_dates degrade to "trial end" due to missing interval info.
    try:
        if paid_flag and subscription_id and _stripe_enabled():
            stripe.api_key = (os.getenv('STRIPE_SECRET_KEY') or '').strip()
            sub_obj = stripe.Subscription.retrieve(subscription_id, expand=["items.data.price"])
            # Always try SubscriptionItem.list to reliably get price/recurring
            # (some Stripe responses omit items.data even with expand).
            si_err = ""
            si_count = 0
            si_price_id = ""
            try:
                si = stripe.SubscriptionItem.list(subscription=subscription_id, limit=10, expand=["data.price"])
                si_data = list(getattr(si, "data", []) or [])
                si_count = len(si_data)
                if si_data:
                    p = getattr(si_data[0], "price", None)
                    if isinstance(p, str):
                        si_price_id = p.strip()
                    elif isinstance(p, dict):
                        si_price_id = str(p.get("id") or "").strip()
                    else:
                        si_price_id = str(getattr(p, "id", "") or "").strip() if p else ""
            except Exception as e:
                si_err = f"{type(e).__name__}: {str(e)}"

            # Stripe can sometimes return plain dicts; handle both dict and StripeObject.
            status = str(_stripe_obj_get(sub_obj, "status", "") or "").strip().lower()
            trial_end = _stripe_obj_get(sub_obj, "trial_end", None)
            current_period_end = _stripe_obj_get(sub_obj, "current_period_end", None)

            # Some Stripe setups can surface an "active" subscription where current_period_end is not populated
            # in our retrieved object. Fallback to upcoming invoice period_end for accurate renewal timing.
            inv_err = ""
            inv_period_end = None
            inv_next_payment_attempt = None
            raw_sub_cpe = None
            raw_sub_billing_anchor = None
            raw_sub_current_period_start = None
            raw_sub_err = ""
            latest_inv_period_end = None
            latest_inv_err = ""
            sanity_applied = False
            sanity_prev_cpe = None
            sanity_new_cpe = None
            if not current_period_end and status in ("active", "trialing"):
                try:
                    inv = _stripe_upcoming_invoice(customer_id, subscription_id)
                    inv_period_end = _stripe_obj_get(inv, "period_end", None)
                    inv_next_payment_attempt = _stripe_obj_get(inv, "next_payment_attempt", None)
                    if inv_period_end:
                        current_period_end = inv_period_end
                except Exception as e:
                    inv_err = f"{type(e).__name__}: {str(e)}"
            # Final fallback: fetch raw subscription JSON and read current_period_end directly.
            if not current_period_end:
                try:
                    raw_sub = _stripe_subscription_raw(subscription_id)
                    raw_sub_cpe = _stripe_obj_get(raw_sub, "current_period_end", None)
                    raw_sub_billing_anchor = _stripe_obj_get(raw_sub, "billing_cycle_anchor", None)
                    raw_sub_current_period_start = _stripe_obj_get(raw_sub, "current_period_start", None)
                    if raw_sub_cpe:
                        current_period_end = raw_sub_cpe
                except Exception:
                    raw_sub_err = "raw_subscription_fetch_failed"

            # Fallback: use latest invoice period_end if available
            if not current_period_end:
                try:
                    latest_inv = _stripe_latest_invoice_for_subscription(subscription_id)
                    latest_inv_period_end = _stripe_obj_get(latest_inv, "period_end", None)
                    if not latest_inv_period_end:
                        # Sometimes invoice line has period.end
                        lines = _stripe_obj_get(latest_inv, "lines", None)
                        ldata = _stripe_obj_get(lines, "data", []) if lines else []
                        if ldata:
                            period = _stripe_obj_get(ldata[0], "period", None)
                            latest_inv_period_end = _stripe_obj_get(period, "end", None) if period else None
                    if latest_inv_period_end:
                        current_period_end = latest_inv_period_end
                except Exception as e:
                    latest_inv_err = f"{type(e).__name__}: {str(e)}"

            # Final fallback: approximate from billing_cycle_anchor/current_period_start + interval.
            if not current_period_end:
                try:
                    base_ts = None
                    if raw_sub_billing_anchor:
                        base_ts = int(raw_sub_billing_anchor)
                    elif raw_sub_current_period_start:
                        base_ts = int(raw_sub_current_period_start)
                    if base_ts:
                        base_dt = datetime.fromtimestamp(base_ts, tz=timezone.utc)
                        if interval:
                            current_period_end = int(_add_interval_approx(base_dt, interval, interval_count).timestamp())
                except Exception:
                    pass

            # Sanity check: sometimes we end up with a "period end" equal to the billing anchor (or not in the future),
            # which makes the UI show today's date. For active subscriptions, ensure period_end is in the future.
            try:
                if status == "active" and interval and current_period_end:
                    now_ts = int(datetime.now(timezone.utc).timestamp())
                    cpe = int(current_period_end)
                    sanity_prev_cpe = cpe
                    anchor_ts = None
                    try:
                        if raw_sub_billing_anchor:
                            anchor_ts = int(raw_sub_billing_anchor)
                        elif raw_sub_current_period_start:
                            anchor_ts = int(raw_sub_current_period_start)
                    except Exception:
                        anchor_ts = None
                    # If cpe is not meaningfully after anchor, or it's already due/expired, recompute.
                    if (anchor_ts is not None and cpe <= (anchor_ts + 60)) or (cpe <= (now_ts + 300)):
                        base_for_calc = anchor_ts if anchor_ts is not None else now_ts
                        base_dt = datetime.fromtimestamp(int(base_for_calc), tz=timezone.utc)
                        current_period_end = int(_add_interval_approx(base_dt, interval, interval_count).timestamp())
                        sanity_applied = True
                        sanity_new_cpe = int(current_period_end)
            except Exception:
                pass

            price_id, interval, interval_count = _get_subscription_price_id_and_recurring(sub_obj)
            if not price_id and si_price_id:
                price_id = si_price_id

            annual_pid = _get_stripe_price_id('annual_6_95') or ''
            monthly_pid = _get_stripe_price_id('monthly_10_95') or ''
            # Infer interval if Stripe didn't provide recurring info
            if not interval and price_id and annual_pid and price_id == annual_pid:
                interval, interval_count = 'year', 1
            if not interval and price_id and monthly_pid and price_id == monthly_pid:
                interval, interval_count = 'month', 1

            # Compute next_billing / paid_through from subscription directly.
            next_ts = None
            if status == "trialing" and trial_end:
                next_ts = int(trial_end)
            elif current_period_end:
                next_ts = int(current_period_end)

            if next_ts:
                next_billing_iso = datetime.fromtimestamp(int(next_ts), tz=timezone.utc).isoformat()

            if status == "trialing" and trial_end:
                base = datetime.fromtimestamp(int(trial_end), tz=timezone.utc)
                if interval:
                    paid_through_est_iso = _add_interval_approx(base, interval, interval_count).isoformat()
                else:
                    # Worst-case fallback: show trial end (better than blank)
                    paid_through_est_iso = base.isoformat()
            elif current_period_end:
                paid_through_est_iso = datetime.fromtimestamp(int(current_period_end), tz=timezone.utc).isoformat()

            if interval == 'year':
                interval_label = interval_label or 'Annual'
                plan_status_raw = plan_status_raw or 'annual_6_95'
            elif interval == 'month':
                interval_label = interval_label or 'Monthly'
                plan_status_raw = plan_status_raw or 'monthly_10_95'

            if debug:
                # also show what Stripe returned around items
                sub_items_len = 0
                try:
                    sub_items_len = len(list(getattr(getattr(sub_obj, "items", None), "data", []) or []))
                except Exception:
                    sub_items_len = 0
                debug_info = {
                    "settings_debug_version": "sanitycheck_v2",
                    "azure_plan_status": str(prof.get("plan_status") or ""),
                    "azure_paid_until": str(prof.get("paid_until") or ""),
                    "stripe_customer_id": customer_id,
                    "stripe_subscription_id": subscription_id,
                    "stripe_status": status,
                    "stripe_trial_end": str(trial_end or ""),
                    "stripe_current_period_end": str(current_period_end or ""),
                    "stripe_upcoming_invoice_period_end": str(inv_period_end or ""),
                    "stripe_upcoming_invoice_next_payment_attempt": str(inv_next_payment_attempt or ""),
                    "stripe_upcoming_invoice_error": inv_err,
                    "stripe_raw_subscription_current_period_end": str(raw_sub_cpe or ""),
                    "stripe_raw_subscription_billing_cycle_anchor": str(raw_sub_billing_anchor or ""),
                    "stripe_raw_subscription_current_period_start": str(raw_sub_current_period_start or ""),
                    "stripe_raw_subscription_error": raw_sub_err,
                    "stripe_latest_invoice_period_end": str(latest_inv_period_end or ""),
                    "stripe_latest_invoice_error": latest_inv_err,
                    "stripe_sanity_applied": str(sanity_applied),
                    "stripe_sanity_prev_cpe": str(sanity_prev_cpe or ""),
                    "stripe_sanity_new_cpe": str(sanity_new_cpe or ""),
                    "stripe_item_price_id": price_id,
                    "stripe_subscription_items_len": str(sub_items_len),
                    "stripe_subscriptionitem_list_count": str(si_count),
                    "stripe_subscriptionitem_list_error": si_err,
                    "stripe_subscriptionitem_price_id": si_price_id,
                    "env_annual_price_id": annual_pid,
                    "env_monthly_price_id": monthly_pid,
                    "computed_interval": interval,
                    "computed_interval_count": str(interval_count),
                    "computed_next_billing_iso": next_billing_iso,
                    "computed_paid_through_iso": paid_through_est_iso,
                }
    except Exception:
        pass

    # If Stripe didn't provide recurring interval details (common with some setups),
    # paid_through_est_iso can end up equal to trial_end. In that case, infer based on our plan_id.
    try:
        if (
            is_paid_user(current_user)
            and next_billing_iso
            and paid_through_est_iso
            and paid_through_est_iso == next_billing_iso
            and plan_status_raw in ('annual_6_95', 'monthly_10_95')
        ):
            trial_end_dt = _parse_iso_dt(next_billing_iso.replace('Z', '+00:00'))
            if trial_end_dt is not None:
                if plan_status_raw == 'annual_6_95':
                    paid_through_est_iso = _add_interval_approx(trial_end_dt, 'year', 1).isoformat()
                    interval_label = interval_label or 'Annual'
                elif plan_status_raw == 'monthly_10_95':
                    paid_through_est_iso = _add_interval_approx(trial_end_dt, 'month', 1).isoformat()
                    interval_label = interval_label or 'Monthly'
    except Exception:
        pass

    # Back-compat: keep paid_until_display populated (prefer Stripe paid-through estimate if it exists)
    if paid_flag:
        paid_until_raw = paid_through_est_iso or paid_until_raw
        if not paid_until_raw:
            # Older fallback: if we only have subscription_id stored, try it.
            sub_id = _get_stripe_subscription_id_from_azure(getattr(current_user, 'id', ''))
            paid_until_raw = _get_paid_until_from_stripe(sub_id) or paid_until_raw
    return render_template(
        'settings.html',
        user=current_user,
        email=(getattr(current_user, 'email', '') or '').strip(),
        name=(getattr(current_user, 'name', '') or '').strip(),
        is_paid=paid_flag,
        plan_status=plan_status_raw,
        interval_label=interval_label,
        next_billing_display=_format_paid_until(next_billing_iso),
        paid_through_display=_format_paid_until(paid_until_raw),
        debug_info=debug_info,
    )

@app.route('/api/me')
def api_me():
    """Lightweight user info for the React frontend (plan gating, limits)."""
    try:
        if not current_user.is_authenticated:
            return jsonify({
                "is_authenticated": False,
                "is_paid": False,
                "free_revision_limit": FREE_REVISION_LIMIT,
                "revisions_used": 0,
            })
        paid = is_paid_user(current_user)
        used = 0
        try:
            used = len(get_user_revisions(current_user.id))
        except Exception:
            used = 0
        return jsonify({
            "is_authenticated": True,
            "is_paid": bool(paid),
            "free_revision_limit": FREE_REVISION_LIMIT,
            "revisions_used": used,
        })
    except Exception:
        return jsonify({
            "is_authenticated": False,
            "is_paid": False,
            "free_revision_limit": FREE_REVISION_LIMIT,
            "revisions_used": 0,
        })

@app.route('/view_revision/<revision_id>')
@login_required
def view_revision(revision_id):
    revisions = get_user_revisions(current_user.id)
    rev = next((r for r in revisions if r['revision_id'] == revision_id), None)
    if not rev:
        flash('Revision not found.', 'danger')
        return redirect(url_for('my_revisions'))

    # Ensure template selection works for saved revisions.
    # The template flow (/api/parse-resume-for-template) reads from session['results_data'].
    # When a user logs out/in, the session is fresh, so we must seed it from the revision
    # being viewed (otherwise template selection fails with "Resume data not found").
    session['results_data'] = {
        'original_resume': rev.get('original_resume', '') or '',
        'revised_resume': rev.get('resume_content', '') or '',
        'feedback': rev.get('feedback', {}) or {},
        'job_description': rev.get('job_description', '') or '',
        'source_revision_id': revision_id,
    }
    # Prevent confusion from a previous template snapshot
    session.pop('template_data', None)
    session.modified = True

    return render_template(
        'result.html',
        revised_resume=rev['resume_content'],
        feedback=rev.get('feedback', {}),
        original_resume=rev.get('original_resume', ''),
        user=current_user
    )

@app.route('/update_notes/<revision_id>', methods=['POST'])
@login_required
def update_notes(revision_id):
    try:
        notes = request.form.get('notes', '').strip()
        
        # Get the current revision to preserve existing data
        table_client = get_table_client()
        entity = table_client.get_entity(partition_key=current_user.id, row_key=revision_id)
        
        # Update only the notes field
        entity['notes'] = notes
        table_client.update_entity(entity)
        
        flash('Notes updated successfully!', 'success')
    except Exception as e:
        flash('Error updating notes. Please try again.', 'danger')
        
    return redirect(url_for('my_revisions'))

@app.route('/application/add/<revision_id>', methods=['POST'])
@login_required
def add_application(revision_id):
    try:
        company = (request.form.get('company') or '').strip()
        role = (request.form.get('role') or '').strip()
        date_applied = (request.form.get('date_applied') or '').strip()
        status = (request.form.get('status') or '').strip()
        link = (request.form.get('link') or '').strip()

        if not any([company, role, date_applied, status, link]):
            flash('Please fill at least one field before adding an application.', 'danger')
            return redirect(url_for('my_revisions'))

        table_client = get_table_client()
        entity = table_client.get_entity(partition_key=current_user.id, row_key=revision_id)
        apps = _parse_applications(entity.get('applications', ''))
        apps.append({
            'company': company,
            'role': role,
            'date_applied': date_applied,
            'status': status,
            'link': link,
        })
        # Prevent unbounded growth
        apps = apps[-200:]
        entity['applications'] = json.dumps(apps)
        table_client.update_entity(entity, mode=UpdateMode.MERGE)
        flash('Application added!', 'success')
    except Exception:
        flash('Error adding application. Please try again.', 'danger')
    return redirect(url_for('my_revisions'))

@app.route('/application/delete/<revision_id>/<int:idx>', methods=['POST'])
@login_required
def delete_application(revision_id, idx: int):
    try:
        table_client = get_table_client()
        entity = table_client.get_entity(partition_key=current_user.id, row_key=revision_id)
        apps = _parse_applications(entity.get('applications', ''))
        if idx < 0 or idx >= len(apps):
            flash('Application not found.', 'danger')
            return redirect(url_for('my_revisions'))
        apps.pop(idx)
        entity['applications'] = json.dumps(apps)
        table_client.update_entity(entity, mode=UpdateMode.MERGE)
        flash('Application removed.', 'success')
    except Exception:
        flash('Error removing application. Please try again.', 'danger')
    return redirect(url_for('my_revisions'))

#############################################
# Admin: Registered Users from Azure Revisions
#############################################

def _collect_registered_users(table_override: str = None):
    """Collect distinct users from Azure Table revisions with counts and basic profile."""
    table_client = get_table_client(table_override)
    counts = {}
    # Ensure we iterate through all pages
    pager = table_client.list_entities(select=["PartitionKey"])
    for entity in pager:
        uid = entity.get("PartitionKey")
        if not uid:
            continue
        counts[uid] = counts.get(uid, 0) + 1

    results = []
    for uid, n in counts.items():
        user_obj = users.get(uid)
        email = getattr(user_obj, "email", "") if user_obj else ""
        name = getattr(user_obj, "name", "") if user_obj else ""
        provider = "google" if str(uid).isdigit() else ("facebook" if str(uid).startswith("facebook_") else "other")
        results.append({
            "id": uid,
            "email": email,
            "name": name,
            "provider": provider,
            "revisions": n,
        })

    results.sort(key=lambda x: x["revisions"], reverse=True)
    return results

@app.route('/admin/registered_users.json')
@login_required
def registered_users_json():
    if not getattr(current_user, "is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    data = _collect_registered_users(request.args.get('table') or None)
    return jsonify({"users": data})

@app.route('/admin/registered_users')
@login_required
def registered_users_view():
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to view this page.", "danger")
        return redirect(url_for("index"))
    data = _collect_registered_users()
    return render_template('admin_registered_users.html', users=data)

@app.route('/admin/registered_users.csv')
@login_required
def registered_users_csv():
    if not getattr(current_user, "is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    rows = _collect_registered_users(request.args.get('table') or None)
    # Build CSV in-memory
    lines = ["id,name,email,provider,revisions"]
    for r in rows:
        # naive CSV escaping for commas and quotes
        def esc(v):
            s = str(v or "")
            if any(c in s for c in [',','"','\n','\r']):
                s = '"' + s.replace('"','""') + '"'
            return s
        line = ",".join([esc(r["id"]), esc(r["name"]), esc(r["email"]), esc(r["provider"]), str(r["revisions"])])
        lines.append(line)
    csv_data = "\n".join(lines) + "\n"
    return Response(csv_data, mimetype='text/csv', headers={
        'Content-Disposition': 'attachment; filename=registered_users.csv'
    })

@app.route('/admin/google_emails.json')
@login_required
def google_emails_json():
    if not getattr(current_user, "is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    # Read directly from persistent users file to include all registered users
    try:
        data = {}
        if os.path.exists(USERS_FILE):
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
        emails = sorted({
            (u.get('email') or '').strip().lower()
            for u in data.values()
            if u and u.get('email') and str(u.get('id', '')).isdigit()
        })
        return jsonify({"count": len(emails), "emails": emails})
    except Exception as e:
        logger.error(f"google_emails_json error: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500

@app.route('/admin/google_emails.csv')
@login_required
def google_emails_csv():
    if not getattr(current_user, "is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    # Read directly from persistent users file to include all registered users
    try:
        data = {}
        if os.path.exists(USERS_FILE):
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
        emails = sorted({
            (u.get('email') or '').strip().lower()
            for u in data.values()
            if u and u.get('email') and str(u.get('id', '')).isdigit()
        })
        csv_data = "email\n" + "\n".join(emails) + "\n"
        return Response(csv_data, mimetype='text/csv', headers={
            'Content-Disposition': 'attachment; filename=google_emails.csv'
        })
    except Exception as e:
        logger.error(f"google_emails_csv error: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500

# Admin: Users from Azure Users table
@app.route('/admin/users_azure.json')
@login_required
def users_azure_json():
    if not getattr(current_user, "is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    provider_filter = (request.args.get('provider') or '').strip().lower()
    try:
        table_client = get_users_table_client()
        users_list = []
        for e in table_client.list_entities():
            if e.get('RowKey') != 'profile':
                continue
            provider = (e.get('provider') or '').lower()
            if provider_filter and provider != provider_filter:
                continue
            users_list.append({
                'id': e.get('PartitionKey'),
                'name': e.get('name') or '',
                'email': e.get('email') or '',
                'provider': provider or 'other',
                'created_at': e.get('created_at') or '',
                'is_admin': bool(e.get('is_admin', False)),
            })
        # Sort by created_at desc if present
        from datetime import datetime
        def parse_dt(s):
            try:
                return datetime.fromisoformat(s)
            except Exception:
                return datetime.min
        users_list.sort(key=lambda x: parse_dt(x.get('created_at') or ''), reverse=True)
        return jsonify({"count": len(users_list), "users": users_list})
    except Exception as e:
        logger.error(f"users_azure_json error: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500

@app.route('/admin/users_azure.csv')
@login_required
def users_azure_csv():
    if not getattr(current_user, "is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    provider_filter = (request.args.get('provider') or '').strip().lower()
    try:
        table_client = get_users_table_client()
        rows = []
        for e in table_client.list_entities():
            if e.get('RowKey') != 'profile':
                continue
            provider = (e.get('provider') or '').lower()
            if provider_filter and provider != provider_filter:
                continue
            rows.append({
                'id': e.get('PartitionKey'),
                'name': e.get('name') or '',
                'email': e.get('email') or '',
                'provider': provider or 'other',
                'created_at': e.get('created_at') or '',
                'is_admin': 'true' if e.get('is_admin', False) else 'false',
            })
        def esc(v):
            s = str(v or "")
            if any(c in s for c in [',','"','\n','\r']):
                s = '"' + s.replace('"','""') + '"'
            return s
        header = ['id','name','email','provider','created_at','is_admin']
        lines = [','.join(header)]
        for r in rows:
            lines.append(','.join([esc(r[h]) for h in header]))
        csv_data = "\n".join(lines) + "\n"
        return Response(csv_data, mimetype='text/csv', headers={
            'Content-Disposition': 'attachment; filename=users_azure.csv'
        })
    except Exception as e:
        logger.error(f"users_azure_csv error: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500

# Admin: set plan/subscription fields on Azure Users profile (manual override until Stripe is integrated)
@app.route('/admin/set_user_plan', methods=['POST'])
@login_required
def admin_set_user_plan():
    if not getattr(current_user, "is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    data = request.get_json(silent=True) or {}
    user_id = str((data.get('user_id') or '')).strip()
    if not user_id:
        return jsonify({"error": "Missing user_id"}), 400
    # Supported fields
    plan_status = str((data.get('plan_status') or '')).strip().lower()  # free|trial|paid|active|canceled
    paid_until = str((data.get('paid_until') or '')).strip()  # ISO timestamp optional
    is_paid_flag = data.get('is_paid', None)  # boolean optional
    try:
        table_client = get_users_table_client()
        # MERGE upsert: do not clobber existing profile fields
        entity = {
            'PartitionKey': user_id,
            'RowKey': 'profile',
        }
        if plan_status:
            entity['plan_status'] = plan_status
        if paid_until:
            entity['paid_until'] = paid_until
        if isinstance(is_paid_flag, bool):
            entity['is_paid'] = bool(is_paid_flag)
        table_client.upsert_entity(entity)
        return jsonify({"status": "ok"}), 200
    except Exception as e:
        logger.error(f"admin_set_user_plan error: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500

def extract_text_from_file(file):
    """Extract text from uploaded file (PDF or DOCX), using fallback for complex Word docs."""
    try:
        filename = secure_filename(file.filename)
        file_extension = filename.lower().split('.')[-1]
        print(f"Processing file: {filename}, extension: {file_extension}")

        if file_extension == 'pdf':
            text = ""
            try:
                with pdfplumber.open(file) as pdf:
                    print(f"PDF has {len(pdf.pages)} pages")
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            print(f"Page {page_num + 1}: {len(page_text)} characters extracted")
            except Exception as e:
                print(f"pdfplumber failed: {str(e)}, trying PyPDF2")
                file.seek(0)
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    print(f"Page {page_num + 1}: {len(page_text)} characters extracted")

            print(f"Total PDF text extracted: {len(text)} characters")
            return text.strip()

        elif file_extension == 'docx':
            
            
            # Try python-docx first
            try:
                doc = Document(file)
                text = ""
                # Create a BytesIO copy for reuse  
                file.seek(0)
                docx_buffer = BytesIO(file.read())
                for para_num, paragraph in enumerate(doc.paragraphs):
                    text += paragraph.text + "\n"
            
                print(f"Total DOC text extracted: {len(text)} characters")

                # Fallback if too little text
                if len(text.strip()) < 75:
                    print("Text too short, falling back to mammoth...")
                    docx_buffer.seek(0)
                    result = mammoth.extract_raw_text(docx_buffer)
                    text = result.value
                    print(f"Text extracted using mammoth: {len(text)} characters")
            except Exception as e:
                print(f"python-docx failed: {str(e)}, trying mammoth as fallback")
                docx_buffer.seek(0)
                result = mammoth.extract_raw_text(docx_buffer)
                text = result.value
                print(f"Text extracted using mammoth: {len(text)} characters")

            return text.strip()

        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    except Exception as e:
        print(f"Error extracting text from file: {str(e)}")
        import traceback
        traceback.print_exc()
        raise ValueError(f"Failed to extract text from file: {str(e)}")

@app.route('/increment_counter', methods=['POST'])
def increment_counter():
    """API endpoint to increment resume counter (conversion tracking)"""
    try:
        # Use analytics module to track conversion
        conversion_result = analytics.track_conversion(session, "resume_submission")
        
        if conversion_result.get("status") == "success":
            # Get updated total conversions count
            analytics_data = analytics.get_full_analytics()
            count = analytics_data.get('summary', {}).get('total_conversions', 0)
            return {"success": True, "count": count}
        else:
            return {"success": False, "error": "Failed to track conversion"}, 500
            
    except Exception as e:
        return {"success": False, "error": str(e)}, 500

@app.route('/counter')
def counter_page():
    """Display counter page"""
    try:
        # Use the analytics module to get the total conversion count
        analytics_data = analytics.get_full_analytics()
        count = analytics_data.get('summary', {}).get('total_conversions', 0)
        
        # Fallback to legacy counter.json format if analytics doesn't work
        if count == 0:
            counter_file = 'counter.json'
            if os.path.exists(counter_file):
                with open(counter_file, 'r') as f:
                    data = json.load(f)
                    # Try new format first (total_conversions), then legacy format (count)
                    count = data.get('total_conversions', data.get('count', 0))
        
        return render_template('counter.html', count=count)
    except Exception as e:
        # Log the error for debugging
        print(f"Counter page error: {str(e)}")
        return render_template('counter.html', count=0)

@app.route('/counter.html')
def counter_html_legacy():
    """Redirect old /counter.html URL to the canonical /counter route."""
    return redirect(url_for('counter_page'), code=301)

@app.route('/robots.txt')
def robots_txt():
    """Serve robots.txt file"""
    try:
        response = send_file('robots.txt', mimetype='text/plain')
        if not app.debug:
            response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
            response.headers['Cache-Control'] = 'public, max-age=86400'  # Cache for 24 hours
        return response
    except Exception as e:
        # Fallback content if file is not found
        content = """User-agent: *
Allow: /
Disallow: /admin/
Disallow: /private/

Sitemap: https://resumaticai.com/sitemap.xml
Host: https://resumaticai.com"""
        response = Response(content, mimetype='text/plain')
        if not app.debug:
            response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
        return response

@app.route('/sitemap-static.xml')
def sitemap_static():
    """Serve static sitemap.xml file as backup"""
    try:
        response = send_file('sitemap.xml', mimetype='application/xml')
        response.headers['Cache-Control'] = 'public, max-age=86400'  # Cache for 24 hours
        return response
    except Exception as e:
        # Fallback to dynamic sitemap
        return redirect(url_for('sitemap'))

@app.route('/ads.txt')
def ads_txt():
    """Serve ads.txt file"""
    response = send_file('ads.txt', mimetype='text/plain')
    if not app.debug:
        response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
        response.headers['Cache-Control'] = 'public, max-age=86400'
    return response

@app.route('/llms.txt')
def llms_txt():
    """Serve llms.txt file"""
    response = send_file('llms.txt', mimetype='text/plain')
    # Add security headers to ensure HTTPS preference
    if not app.debug:
        response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
        response.headers['Cache-Control'] = 'public, max-age=86400'
    return response

@app.route('/subscribers')
@login_required
def subscribers():
    """Show subscribers page with actual subscriber data"""
    # Check if the user is an admin
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to view this page.", "danger")
        return redirect(url_for("index"))
    
    try:
        subscriber_list = []
        subscriber_count = 0
        
        # Read subscribers from CSV file
        if os.path.exists("subscribers.csv"):
            with open("subscribers.csv", "r") as f:
                lines = f.readlines()
                # Skip header row if present
                for line in lines[1:] if lines and lines[0].strip().lower() == 'email' else lines:
                    email = line.strip()
                    if email:  # Only add non-empty emails
                        subscriber_list.append({
                            'email': email,
                            'date_added': 'Unknown'  # CSV doesn't store dates
                        })
                        subscriber_count += 1
        
        # Sort subscribers alphabetically
        subscriber_list.sort(key=lambda x: x['email'].lower())
        
        return render_template('subscribers.html', 
                             subscribers=subscriber_list, 
                             total_count=subscriber_count,
                             user=current_user)
    except Exception as e:
        flash(f"Error loading subscribers: {str(e)}", "danger")
        return render_template('subscribers.html', 
                             subscribers=[], 
                             total_count=0,
                             user=current_user)

@app.route('/signup')
def signup():
    """Show signup page"""
    return render_template('signup.html')

# Newsletter Management Routes
@app.route('/admin/newsletter')
@login_required
def newsletter_admin():
    """Newsletter administration dashboard"""
    if not current_user.is_authenticated:
        flash("You need to log in to view this page.", "danger")
        return redirect(url_for("login"))

    # Check if the user is an admin
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to view this page.", "danger")
        return redirect(url_for("index"))

    try:
        # Get newsletter archives
        newsletter_files = []
        if os.path.exists("newsletters"):
            for filename in os.listdir("newsletters"):
                if filename.endswith(".json"):
                    try:
                        with open(os.path.join("newsletters", filename), "r") as f:
                            data = json.load(f)
                        newsletter_files.append({
                            "filename": filename,
                            "month": data.get("month"),
                            "year": data.get("year"),
                            "subject": data.get("subject"),
                            "generated_at": data.get("generated_at"),
                            "send_stats": data.get("send_stats", {})
                        })
                    except Exception as e:
                        logger.error(f"Error reading newsletter file {filename}: {str(e)}")

        # Sort by year and month (newest first)
        newsletter_files.sort(key=lambda x: (x.get("year", 0), x.get("month", "")), reverse=True)

        # Get subscriber count
        subscriber_count = 0
        if os.path.exists("subscribers.csv"):
            with open("subscribers.csv", "r") as f:
                lines = f.readlines()
                subscriber_count = len([line for line in lines[1:] if line.strip()]) if lines else 0

        return render_template('newsletter_admin.html', 
                             newsletters=newsletter_files, 
                             subscriber_count=subscriber_count,
                             user=current_user)
    except Exception as e:
        flash(f"Error loading newsletter dashboard: {str(e)}", "danger")
        return redirect(url_for("admin_stats"))

@app.route('/admin/newsletter/generate', methods=['POST'])
@login_required
def generate_newsletter():
    """Generate a new newsletter"""
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to access this feature.", "danger")
        return redirect(url_for("index"))

    try:
        # Get form data
        month = request.form.get('month')
        year_str = request.form.get('year')
        custom_topics_str = request.form.get('custom_topics', '').strip()
        preview_only = request.form.get('preview_only') == 'on'

        # Validate inputs
        if not month or not year_str:
            flash("Please select both month and year.", "danger")
            return redirect(url_for('newsletter_admin'))

        try:
            year = int(year_str)
        except ValueError:
            flash("Invalid year provided.", "danger")
            return redirect(url_for('newsletter_admin'))

        # Parse custom topics
        custom_topics = None
        if custom_topics_str:
            custom_topics = [topic.strip() for topic in custom_topics_str.split(',') if topic.strip()]

        # Initialize newsletter manager
        newsletter_manager = NewsletterManager()

        # Generate newsletter
        result = newsletter_manager.create_and_send_newsletter(
            month=month,
            year=year,
            custom_topics=custom_topics,
            preview_only=preview_only
        )

        if preview_only:
            flash(f"Newsletter preview generated for {month} {year}! Check the archives to view it.", "success")
        else:
            send_stats = result.get('send_stats', {})
            flash(f"Newsletter sent! {send_stats.get('sent', 0)} emails sent successfully, {send_stats.get('failed', 0)} failed.", "success")

        return redirect(url_for('newsletter_admin'))

    except Exception as e:
        logger.error(f"Error generating newsletter: {str(e)}")
        flash(f"Error generating newsletter: {str(e)}", "danger")
        return redirect(url_for('newsletter_admin'))

@app.route('/admin/newsletter/preview/<path:filename>')
@login_required
def preview_newsletter(filename):
    """Preview a newsletter"""
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to access this feature.", "danger")
        return redirect(url_for("index"))

    try:
        # Security check - ensure filename is safe
        safe_filename = secure_filename(filename)
        if not safe_filename.endswith('.html'):
            safe_filename += '.html'

        filepath = os.path.join("newsletters", safe_filename)
        
        if not os.path.exists(filepath):
            flash("Newsletter not found.", "danger")
            return redirect(url_for('newsletter_admin'))

        return send_file(filepath)

    except Exception as e:
        logger.error(f"Error previewing newsletter: {str(e)}")
        flash("Error loading newsletter preview.", "danger")
        return redirect(url_for('newsletter_admin'))

@app.route('/admin/newsletter/send/<path:filename>', methods=['POST'])
@login_required
def send_existing_newsletter(filename):
    """Send an existing newsletter"""
    if not getattr(current_user, "is_admin", False):
        flash("You do not have permission to access this feature.", "danger")
        return redirect(url_for("index"))

    try:
        # Security check
        safe_filename = secure_filename(filename)
        if not safe_filename.endswith('.json'):
            safe_filename += '.json'

        filepath = os.path.join("newsletters", safe_filename)
        
        if not os.path.exists(filepath):
            flash("Newsletter not found.", "danger")
            return redirect(url_for('newsletter_admin'))

        # Load newsletter data
        with open(filepath, "r") as f:
            newsletter_data = json.load(f)

        # Initialize newsletter manager and send
        newsletter_manager = NewsletterManager()
        send_stats = newsletter_manager.sender.send_newsletter(
            subject=newsletter_data["subject"],
            html_content=newsletter_data["html_content"],
            text_content=newsletter_data["text_content"]
        )

        # Update the newsletter file with send stats
        newsletter_data["send_stats"] = send_stats
        newsletter_data["last_sent"] = datetime.now(timezone.utc).isoformat()
        
        with open(filepath, "w") as f:
            json.dump(newsletter_data, f, indent=2)

        flash(f"Newsletter sent! {send_stats.get('sent', 0)} emails sent successfully, {send_stats.get('failed', 0)} failed.", "success")

    except Exception as e:
        logger.error(f"Error sending newsletter: {str(e)}")
        flash(f"Error sending newsletter: {str(e)}", "danger")

    return redirect(url_for('newsletter_admin'))

@app.route('/api/newsletter/test', methods=['POST'])
@login_required
def test_newsletter_config():
    """Test newsletter configuration (SMTP settings)"""
    if not getattr(current_user, "is_admin", False):
        return jsonify({"status": "error", "message": "Unauthorized"}), 403

    try:
        config = NewsletterConfig()
        
        # Check if credentials are configured
        if not config.sender_email or not config.sender_password:
            return jsonify({
                "status": "error", 
                "message": "Email credentials not configured. Please add NEWSLETTER_EMAIL and NEWSLETTER_PASSWORD to your .env file."
            }), 400

        # Log the configuration for debugging (without password)
        logger.info(f"Testing email config - Email: {config.sender_email}, SMTP: {config.smtp_server}:{config.smtp_port}")

        # Test SMTP connection
        import smtplib
        server = smtplib.SMTP(config.smtp_server, config.smtp_port)
        server.starttls()
        server.login(config.sender_email, config.sender_password)
        server.quit()

        return jsonify({
            "status": "success", 
            "message": f"Email configuration is working correctly! Connected to {config.smtp_server} with {config.sender_email}"
        })

    except smtplib.SMTPAuthenticationError as e:
        error_msg = str(e)
        if "Username and Password not accepted" in error_msg:
            return jsonify({
                "status": "error", 
                "message": "Authentication failed. For Gmail, make sure you're using an App Password, not your regular password. See newsletter_config.txt for setup instructions."
            }), 400
        else:
            return jsonify({
                "status": "error", 
                "message": f"SMTP Authentication Error: {error_msg}"
            }), 400
    except smtplib.SMTPException as e:
        return jsonify({
            "status": "error", 
            "message": f"SMTP Error: {str(e)}"
        }), 400
    except Exception as e:
        logger.error(f"Newsletter config test error: {str(e)}")
        return jsonify({
            "status": "error", 
            "message": f"Configuration error: {str(e)}"
        }), 400

@app.route('/api/newsletter/debug', methods=['POST'])
@login_required
def debug_newsletter_config():
    """Debug newsletter configuration"""
    if not getattr(current_user, "is_admin", False):
        return jsonify({"status": "error", "message": "Unauthorized"}), 403

    try:
        config = NewsletterConfig()
        
        debug_info = {
            "email_configured": bool(config.sender_email),
            "password_configured": bool(config.sender_password),
            "email_value": config.sender_email if config.sender_email else "Not set",
            "smtp_server": config.smtp_server,
            "smtp_port": config.smtp_port,
            "env_vars": {
                "NEWSLETTER_EMAIL": "Set" if os.getenv("NEWSLETTER_EMAIL") else "Not set",
                "NEWSLETTER_PASSWORD": "Set" if os.getenv("NEWSLETTER_PASSWORD") else "Not set"
            }
        }
        
        return jsonify({"status": "success", "debug_info": debug_info})

    except Exception as e:
        return jsonify({"status": "error", "message": f"Debug error: {str(e)}"}), 400

@app.route('/unsubscribe')
def unsubscribe():
    """Unsubscribe page for newsletter"""
    return render_template('unsubscribe.html')

@app.route('/unsubscribe', methods=['POST'])
def unsubscribe_post():
    """Process unsubscribe request"""
    email = request.form.get('email', '').strip().lower()
    
    if not email:
        flash("Please enter your email address.", "danger")
        return redirect(url_for('unsubscribe'))
    
    try:
        # Read current subscribers
        subscribers = []
        if os.path.exists("subscribers.csv"):
            with open("subscribers.csv", "r") as f:
                subscribers = [line.strip().lower() for line in f.readlines()]
        
        # Check if email exists
        if email not in subscribers:
            flash("Email address not found in our subscriber list.", "info")
            return redirect(url_for('unsubscribe'))
        
        # Remove email from list
        updated_subscribers = [sub for sub in subscribers if sub != email and sub != 'email']
        
        # Write back to file
        with open("subscribers.csv", "w") as f:
            f.write("email\n")  # Header
            for subscriber in updated_subscribers:
                if subscriber:  # Skip empty lines
                    f.write(subscriber + "\n")
        
        flash("You have been successfully unsubscribed from our newsletter.", "success")
        
    except Exception as e:
        logger.error(f"Error unsubscribing email: {str(e)}")
        flash("An error occurred. Please try again later.", "danger")
    
    return redirect(url_for('unsubscribe'))



    # Access all users
for user_id, user_obj in users.items():
    print(f"User: {user_obj.name}, Email: {user_obj.email}")


def _load_email_config_if_missing() -> None:
    """Load/override SMTP creds from newsletter_config.txt into environment."""
    import os
    try:
        if os.path.exists('newsletter_config.txt'):
            with open('newsletter_config.txt', 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('#'):
                        continue
                    if '=' in line:
                        key, val = line.split('=', 1)
                        key = key.strip()
                        val = val.strip()
                        if key in ('NEWSLETTER_EMAIL', 'NEWSLETTER_PASSWORD', 'SMTP_SERVER', 'SMTP_PORT', 'CONTACT_RECIPIENT') and val:
                            os.environ[key] = val
    except Exception as e:
        logger.error(f"Failed to load newsletter_config.txt: {str(e)}")

# Contact form email helper
def send_contact_email(name: str, sender_email: str, message: str) -> None:
    """Send contact form submission to the site owner via SMTP.

    Uses NEWSLETTER_EMAIL/NEWSLETTER_PASSWORD for SMTP auth to avoid duplicating config.
    Recipient defaults to CONTACT_RECIPIENT or the site owner's email.
    """
    import os
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    _load_email_config_if_missing()
    smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
    smtp_port = int(os.getenv('SMTP_PORT', '587'))
    auth_email = os.getenv('NEWSLETTER_EMAIL', '').strip()
    auth_password = os.getenv('NEWSLETTER_PASSWORD', '').strip().replace(' ', '')
    recipient = os.getenv('CONTACT_RECIPIENT', 'yaronyaronlid@gmail.com').strip()

    if not auth_email or not auth_password:
        raise ValueError('Email credentials not configured. Set NEWSLETTER_EMAIL and NEWSLETTER_PASSWORD.')

    subject = 'New Contact Form Submission - ResumaticAI'
    text_body = (
        f"You have a new contact form submission from ResumaticAI.\n\n"
        f"Name: {name or 'N/A'}\n"
        f"Email: {sender_email or 'N/A'}\n\n"
        f"Message:\n{message or ''}\n"
    )

    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = auth_email
    msg['To'] = recipient
    # Prefer plain text; add Reply-To for easy response
    msg.add_header('Reply-To', sender_email or auth_email)
    msg.attach(MIMEText(text_body, 'plain', 'utf-8'))

    server = smtplib.SMTP(smtp_server, smtp_port)
    server.starttls()
    server.login(auth_email, auth_password)
    server.send_message(msg)
    server.quit()


def save_contact_message(name: str, sender_email: str, message: str) -> None:
    """Persist contact messages locally if email delivery fails."""
    import csv
    from datetime import datetime
    filename = 'contact_messages.csv'
    try:
        file_exists = os.path.exists(filename)
        with open(filename, 'a', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if not file_exists:
                writer.writerow(['timestamp_iso', 'name', 'email', 'message'])
            writer.writerow([datetime.utcnow().isoformat(), name or '', sender_email or '', message or ''])
    except Exception as e:
        logger.error(f'Failed to save contact message fallback: {str(e)}')

# Contact form route
@app.route('/contact', methods=['GET', 'POST'])
def contact():
    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        message = request.form.get('message')
        website = request.form.get('website', '')  # honeypot
        try:
            # Basic spam checks: honeypot must be empty, submission must not be too fast (<2s)
            ts_str = request.form.get('_ts', '0')
            try:
                form_ts = int(ts_str)
            except ValueError:
                form_ts = 0
            import time
            now_s = int(time.time())
            # Lower threshold to reduce false positives from autofill
            too_fast = (now_s - form_ts) < 1 if form_ts else False

            if website.strip() or too_fast:
                logger.info('Contact form blocked by spam checks (honeypot/timing).')
                flash('Thank you for contacting us!', 'success')
                return render_template('contact.html', form_ts=int(time.time()))

            send_contact_email(name, email, message)
            flash('Thank you for contacting us! Your message has been sent.', 'success')
        except Exception as e:
            logger.error(f"Contact form email failed: {str(e)}")
            # Fallback: persist the message so it's not lost
            try:
                save_contact_message(name, email, message)
                flash('Thank you for contacting us! We received your message.', 'info')
            except Exception:
                flash('We could not send your message due to a server error. Please try again later.', 'danger')
        return render_template('contact.html', form_ts=int(time.time()))
    # GET: set initial timestamp
    import time
    return render_template('contact.html', form_ts=int(time.time()))

# Offline page route
@app.route('/offline.html')
def offline():
    """Offline page for PWA functionality"""
    return render_template('offline.html')

# 410 Gone for legacy/old thin URLs
@app.route('/index_old')
@app.route('/index_old.html')
@app.route('/indexfiver')
@app.route('/indexfiver.html')
@app.route('/templates.html')
def gone_legacy_urls():
    return Response('This URL has been permanently removed.', status=410, mimetype='text/plain')

# Minimal site search endpoint to support Sitelinks SearchBox
@app.route('/search')
def site_search():
    query = request.args.get('q', '').strip()
    # Simple, non-indexable helper page for users and bots
    html = f"""<!DOCTYPE html>
<html lang=\"en\"><head><meta charset=\"utf-8\">\n<meta name=\"robots\" content=\"noindex, nofollow\">\n<title>Search | ResumaticAI</title></head>
<body style=\"font-family: system-ui, -apple-system, Segoe UI, Roboto, sans-serif; padding: 24px;\">\n<h1>Search</h1>\n<p>Showing results for: <strong>{query}</strong></p>\n<p>We don't have on-site search yet. Try exploring our <a href=\"/blog\">blog</a> or the <a href=\"/\">homepage</a>.</p>\n</body></html>"""
    return Response(html, mimetype='text/html')

# Explicitly set X-Robots-Tag for low-value or transactional routes
@app.after_request
def add_robots_headers(response):
    try:
        noindex_endpoints = {
            'results_route',
            'login',
            'signup',
            'thank_you',
            'unsubscribe',
            'counter_page',
            'counter_html_legacy',
        }
        noindex_paths = {
            '/results',
            '/login',
            '/signup',
            '/thank_you',
            '/unsubscribe',
            '/counter',
            '/counter.html',
        }
        if (request.endpoint in noindex_endpoints) or (request.path in noindex_paths):
            response.headers['X-Robots-Tag'] = 'noindex, nofollow'
    except Exception:
        pass
    return response




#########################################################
#################TEMPLATES
'''
# --- imports (single set) ---
from flask import Flask, render_template, request, send_file, abort, url_for, redirect, jsonify
from io import BytesIO
import os



# Optional: nicer DOCX output if available
try:
    from docx import Document
except Exception:
    Document = None

# ---- Template registry (IDs match the gallery data) ----
TEMPLATES = {
    "modern-citrus": {"name": "Modern Citrus", "style": "Modern"},
    "modern-slate": {"name": "Modern Slate", "style": "Modern"},
    "modern-aurora": {"name": "Modern Aurora", "style": "Modern"},
    "classic-elegant": {"name": "Classic Elegant", "style": "Classic"},
    "classic-simplicity": {"name": "Classic Simplicity", "style": "Classic"},
    "classic-structure": {"name": "Classic Structure", "style": "Classic"},
    "creative-pastel": {"name": "Creative Pastel", "style": "Creative"},
    "creative-neo": {"name": "Creative Neo", "style": "Creative"},
    "creative-spark": {"name": "Creative Spark", "style": "Creative"},
}

def get_template_or_404(tid: str):
    t = TEMPLATES.get(tid)
    if not t:
        abort(404, f"Unknown template id: {tid}")
    return t

# ---- Pages ----
@app.route("/resume-templates", endpoint="resume_templates")
def resume_templates_view():
    return render_template("resume_templates.html")

@app.route("/builder", endpoint="resume_builder")
def resume_builder():
    tpl = request.args.get("template", "")
    t = TEMPLATES.get(tpl)
    return render_template("builder.html", template_id=tpl, template_meta=t)

# Optional alias if old links use /app/builder
@app.route("/app/builder", endpoint="resume_builder_alias")
def resume_builder_alias():
    qs = request.query_string.decode("utf-8")
    target = url_for("resume_builder")
    return redirect(f"{target}?{qs}" if qs else target, code=302)

# ---- Minimal export endpoints ----
@app.post("/export/docx")
def export_docx():
    data = request.get_json(force=True)
    tpl_id = data.get("template")
    t = get_template_or_404(tpl_id)

    name = data.get("name", "")
    title = data.get("title", "")
    summary = data.get("summary", "")
    exp = data.get("experience", []) or []
    skills = data.get("skills", []) or []

    bio = BytesIO()
    if Document:
        doc = Document()
        doc.add_heading(name or "Your Name", 0)
        if title:
            doc.add_paragraph(title)
        doc.add_paragraph(f"Template: {t['name']} ({t['style']})")
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary or "")
        doc.add_heading("Experience", level=1)
        for item in exp:
            p = doc.add_paragraph()
            p.add_run(item.get("role", "Role")).bold = True
            company = item.get("company", "Company")
            years = item.get("years", "")
            p.add_run(f" — {company}" + (f" ({years})" if years else ""))
            for b in item.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")
        if skills:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(skills))
        doc.save(bio)
    else:
        bio.write(
            f"{name}\n{title}\nTemplate: {t['name']} ({t['style']})\n\nSummary:\n{summary}\n".encode("utf-8")
        )
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"{tpl_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

@app.post("/export/pdf")
def export_pdf():
    # Tiny placeholder PDF. For production: WeasyPrint / wkhtmltopdf / ReportLab.
    data = request.get_json(force=True)
    tpl_id = data.get("template")
    t = get_template_or_404(tpl_id)
    name = data.get("name", "")
    title = data.get("title", "")
    text = f"{name} — {title} | {t['name']} ({t['style']})".replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    buf = BytesIO()
    buf.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    objs = []
    objs.append("1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    objs.append("2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    objs.append("3 0 obj\n<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 5 0 R >> >> /MediaBox [0 0 612 792] /Contents 4 0 R >>\nendobj\n")
    stream = f"BT /F1 18 Tf 72 720 Td ({text}) Tj ET"
    objs.append(f"4 0 obj\n<< /Length {len(stream)} >>\nstream\n{stream}\nendstream\nendobj\n")
    objs.append("5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")

    offsets = []
    for o in objs:
        offsets.append(buf.tell())
        buf.write(o.encode("utf-8"))
    xref = buf.tell()
    buf.write(f"xref\n0 {len(objs)+1}\n".encode("utf-8"))
    buf.write(b"0000000000 65535 f \n")
    for off in offsets:
        buf.write(f"{off:010d} 00000 n \n".encode("utf-8"))
    buf.write(f"trailer\n<< /Size {len(objs)+1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF".encode("utf-8"))
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name=f"{tpl_id}.pdf", mimetype="application/pdf")

from jinja2 import TemplateNotFound
import os

@app.route("/", endpoint="home")
def home():
    # Serve index.html if present; otherwise fall back to the gallery
    try:
        return render_template("index.html")
    except TemplateNotFound:
        return redirect(url_for("resume_templates"))

##########################################
'''
from flask import send_from_directory
import os

import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FILE_PATH = os.path.join(BASE_DIR, '610528eb-2434-4e0c-b4c5-1f54f21877ea.html')


@app.route('/610528eb-2434-4e0c-b4c5-1f54f21877ea.html')
def serve_file():
    return send_from_directory(
        directory=BASE_DIR,
        path='610528eb-2434-4e0c-b4c5-1f54f21877ea.html'
    )




import os
from flask import send_file

@app.route("/react")
@app.route("/react/<path:subpath>")
def react_app(subpath=None):
    return send_file(os.path.join("static", "react", "index.html"))




# Schema definitions
RESUME_STRUCTURE_SCHEMA = """
{
    "name": "string",
    "email": "string",
    "phone": "string",
    "location": "string",
    "summary": "string",
    "experience": [
        {
            "title": "string",
            "company": "string",
            "duration": "string",
            "description": "string (bullet points as newline separated text)"
        }
    ],
    "education": [
        {
            "degree": "string",
            "institution": "string",
            "year": "string"
        }
    ],
    "projects": [
        {
            "title": "string",
            "description": "string",
            "technologies": "string (comma separated)",
            "link": "string (optional)"
        }
    ],
    "certifications": [
        {
            "name": "string",
            "issuer": "string",
            "year": "string"
        }
    ],
    "skills": ["string"],
    "custom_sections": [
        {
            "heading": "string (The original section title, e.g. 'Publications', 'Volunteering', 'Awards')",
            "items": [
                {
                    "title": "string (e.g. Award Name or Role)",
                    "subtitle": "string (e.g. Organization or Event)",
                    "date": "string (optional)",
                    "content": "string (Description or details)"
                }
            ]
        }
    ]
}
"""


def parse_resume(revised_resume):
    try:
        # Initialize the client inside the function to avoid blocking startup
        # Explicitly pass API key to handle Azure environment
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set")
        
        # Configure timeout and other parameters for Azure reliability
        client = OpenAI(
            api_key=api_key,
            timeout=60.0,  # 60 second timeout
            max_retries=3  # Retry up to 3 times on transient errors
        )
        
        parsing_prompt = f"""
You are a data extraction engine. 
Extract structured data from the following RESUME TEXT into the specified JSON format.

Resume Text:
\"\"\"{revised_resume}\"\"\"

Output Format:
{RESUME_STRUCTURE_SCHEMA}
        
Rules:
1. Extract the content exactly as written in the source text.
2. Map standard sections to standard fields.
3. Map non-standard sections to 'custom_sections'.
4. Return ONLY valid JSON.
"""
        parsing_response = client.chat.completions.create(
            model="gpt-4o", 
            messages=[
                {"role": "system", "content": "You are a data extraction engine. Output only valid JSON."},
                {"role": "user", "content": parsing_prompt}
            ],
            response_format={"type": "json_object"}
        )

        parsing_content = parsing_response.choices[0].message.content
        
        # Remove markdown backticks if present (similar to revise_resume function)
        if parsing_content.startswith("```"):
            import re
            parsing_content = re.sub(r"^```(?:json)?\n", "", parsing_content)
            parsing_content = re.sub(r"\n```$", "", parsing_content)
            
        try:
            structured_resume = json.loads(parsing_content)
        except json.JSONDecodeError as e:
            logging.error(f"JSON Parse Error in parse_resume: {str(e)}")
            logging.error(f"Raw parsing content: {parsing_content[:500]}...")
            raise ValueError(f"Failed to parse resume structure: {str(e)}")
        return {"resume": structured_resume}
    except Exception as e:
        logging.error(f"Error in parse_resume: {str(e)}")
        import traceback
        traceback.print_exc()
        raise
           


if __name__ == "__main__":
    app.run(debug=True, host='127.0.0.1', port=5000)

